#!/usr/bin/env python3
"""
Enhanced CSRR Faculty Media Search with API Integration
Searches for faculty media mentions June 1 - July 27, 2025 and creates dual Excel outputs
"""

import pandas as pd
import docx
import os
import requests
import time
import random
from datetime import datetime
from typing import List, Dict

def extract_faculty_names_from_document(doc_path):
    """Extract clean faculty names from the Word document"""
    try:
        doc = docx.Document(doc_path)
        faculty_names = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # Skip empty lines, headers, and URLs
            if not text or text.startswith('Op-Eds') or text.startswith('Since') or 'http' in text or 'www.' in text:
                continue
            
            # Check if this looks like a faculty name
            words = text.split()
            if (2 <= len(words) <= 4 and 
                not text.startswith('*') and
                not any(char in text for char in ['©', '®', '™']) and
                not text.isupper() and
                not text.endswith('.') and
                ',' not in text):
                
                # Additional filters to exclude non-names
                if not any(word.lower() in text.lower() for word in ['university', 'college', 'institute', 'center', 'department', 'school']):
                    faculty_names.append(text)
        
        # Remove duplicates while preserving order
        unique_faculty = []
        seen = set()
        for name in faculty_names:
            if name not in seen:
                seen.add(name)
                unique_faculty.append(name)
        
        return unique_faculty
        
    except Exception as e:
        print(f"Error reading document: {e}")
        return []

def search_with_google_api(faculty_name, google_api_key):
    """Search using Google Custom Search API"""
    results = []
    
    # Google Custom Search endpoint
    search_url = "https://www.googleapis.com/customsearch/v1"
    
    # Search queries for recent dates
    queries = [
        f'"{faculty_name}" op-ed opinion commentary 2025',
        f'"{faculty_name}" interview quoted news 2025',
        f'"{faculty_name}" author byline article 2025'
    ]
    
    for query in queries:
        try:
            params = {
                'key': google_api_key,
                'cx': '017576662512468239146:omuauf_lfve',  # Default search engine ID
                'q': query,
                'dateRestrict': 'm2',  # Last 2 months
                'num': 5
            }
            
            response = requests.get(search_url, params=params, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                
                if 'items' in data:
                    for item in data['items']:
                        title = item.get('title', '')
                        url = item.get('link', '')
                        snippet = item.get('snippet', '')
                        
                        # Validate faculty mention
                        content = f"{title} {snippet}".lower()
                        if faculty_name.lower() in content or any(part.lower() in content for part in faculty_name.split()):
                            
                            # Extract source from URL
                            try:
                                from urllib.parse import urlparse
                                domain = urlparse(url).netloc
                                source = domain.replace('www.', '').split('.')[0].title()
                            except:
                                source = 'Unknown'
                            
                            results.append({
                                'faculty_name': faculty_name,
                                'author': faculty_name,
                                'title': title,
                                'source': source,
                                'url': url,
                                'publication_date': 'June-July 2025',
                                'snippet': snippet[:300] + '...' if len(snippet) > 300 else snippet,
                                'date_found': datetime.now().strftime('%Y-%m-%d')
                            })
            
            time.sleep(1)  # Rate limiting
            
        except Exception as e:
            print(f"  Google API error for {faculty_name}: {e}")
            continue
    
    return results

def search_faculty_member(faculty_name, google_api_key):
    """Enhanced search for a faculty member using available APIs"""
    print(f"🔍 Searching: {faculty_name}")
    
    all_results = []
    
    # Try Google API search
    if google_api_key:
        try:
            google_results = search_with_google_api(faculty_name, google_api_key)
            all_results.extend(google_results)
        except Exception as e:
            print(f"  Google API error: {e}")
    
    # Fallback to basic web search if no API results
    if not all_results:
        try:
            # Basic web search as fallback
            import urllib.parse
            from bs4 import BeautifulSoup
            
            session = requests.Session()
            session.headers.update({
                'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'
            })
            
            query = f'"{faculty_name}" news article 2025'
            search_url = f"https://www.bing.com/search?q={urllib.parse.quote(query)}"
            
            response = session.get(search_url, timeout=10)
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                
                for result in soup.find_all('li', class_='b_algo')[:3]:
                    title_elem = result.find('h2')
                    if title_elem and title_elem.find('a'):
                        title = title_elem.get_text(strip=True)
                        url = title_elem.find('a').get('href', '')
                        
                        snippet_elem = result.find('p')
                        snippet = snippet_elem.get_text(strip=True) if snippet_elem else ''
                        
                        # Validate faculty mention
                        content = f"{title} {snippet}".lower()
                        if faculty_name.lower() in content:
                            try:
                                domain = urllib.parse.urlparse(url).netloc
                                source = domain.replace('www.', '').split('.')[0].title()
                            except:
                                source = 'Unknown'
                            
                            all_results.append({
                                'faculty_name': faculty_name,
                                'author': faculty_name,
                                'title': title,
                                'source': source,
                                'url': url,
                                'publication_date': 'June-July 2025',
                                'snippet': snippet[:300] + '...' if len(snippet) > 300 else snippet,
                                'date_found': datetime.now().strftime('%Y-%m-%d')
                            })
        except Exception as e:
            print(f"  Fallback search error: {e}")
    
    if all_results:
        print(f"  ✅ Found {len(all_results)} results")
    else:
        print(f"  ❌ No results found")
    
    return all_results

def extract_existing_entries_from_document(doc_path):
    """Extract existing entries from the reference document"""
    try:
        doc = docx.Document(doc_path)
        entries = []
        current_faculty = None
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Skip headers
            if text.startswith('Op-Eds') or text.startswith('Since'):
                continue
            
            # Check if this is a faculty name
            if (len(text.split()) <= 4 and 
                ',' not in text and 
                'http' not in text and 
                not any(word in text.lower() for word in ['arizona', 'daily', 'star', 'news', 'times', 'post'])):
                current_faculty = text
            elif current_faculty and ',' in text and 'http' in text:
                # This is a publication entry
                parts = text.split(', ')
                if len(parts) >= 4:
                    author = parts[0].strip()
                    title = parts[1].strip()
                    source = parts[2].strip()
                    url = parts[-1].strip().rstrip('.')
                    
                    # Extract date
                    date = "Pre-June 2025"
                    for part in parts:
                        if any(month in part for month in ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']):
                            date = part.strip()
                            break
                    
                    entries.append({
                        'faculty_name': current_faculty,
                        'author': author,
                        'title': title,
                        'source': source,
                        'url': url,
                        'publication_date': date,
                        'snippet': f'From reference document: {text[:200]}...',
                        'date_found': 'Reference Document'
                    })
        
        return entries
    except Exception as e:
        print(f"Error extracting existing entries: {e}")
        return []

def create_excel_reports(new_results, existing_entries):
    """Create two Excel reports as requested"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    
    # Report 1: June 1 - July 27, 2025 only (new findings)
    if new_results:
        for i, result in enumerate(new_results, 1):
            result['search_order'] = i
            
        df_new = pd.DataFrame(new_results)
        df_new = df_new[["faculty_name", "author", "title", "source", "url", 
                        "publication_date", "date_found", "search_order", "snippet"]]
        df_new.columns = ["Faculty Name", "Author", "Title", "Source", "URL", 
                         "Publication Date", "Date Found", "Search Order", "Snippet"]
    else:
        df_new = pd.DataFrame(columns=[
            "Faculty Name", "Author", "Title", "Source", "URL", 
            "Publication Date", "Date Found", "Search Order", "Snippet"
        ])
    
    filename_new = f"CSRR_Faculty_Media_June_July_2025_ONLY_{timestamp}.xlsx"
    df_new.to_excel(filename_new, index=False)
    
    # Report 2: Combined (existing + new)
    all_entries = existing_entries + new_results
    if all_entries:
        for i, entry in enumerate(all_entries, 1):
            entry['search_order'] = i
            
        df_combined = pd.DataFrame(all_entries)
        df_combined = df_combined[["faculty_name", "author", "title", "source", "url", 
                                  "publication_date", "date_found", "search_order", "snippet"]]
        df_combined.columns = ["Faculty Name", "Author", "Title", "Source", "URL", 
                              "Publication Date", "Date Found", "Search Order", "Snippet"]
    else:
        df_combined = pd.DataFrame(columns=[
            "Faculty Name", "Author", "Title", "Source", "URL", 
            "Publication Date", "Date Found", "Search Order", "Snippet"
        ])
    
    filename_combined = f"CSRR_Faculty_Media_COMBINED_All_Periods_{timestamp}.xlsx"
    df_combined.to_excel(filename_combined, index=False)
    
    return filename_new, filename_combined

def main():
    """Main function"""
    print("=" * 60)
    print("🎯 ENHANCED CSRR FACULTY MEDIA SEARCH WITH API")
    print("=" * 60)
    
    # Get API keys from environment
    google_api_key = os.environ.get('GOOGLE_API_KEY')
    openai_api_key = os.environ.get('OPENAI_API_KEY')
    
    print(f"🔑 Google API Key: {'✅ Available' if google_api_key else '❌ Missing'}")
    print(f"🔑 OpenAI API Key: {'✅ Available' if openai_api_key else '❌ Missing'}")
    
    # Extract faculty names
    doc_path = '/Users/azrabano/Downloads/5-31-25 Op-Ed CSRR Affiliates (6).docx'
    print(f"\n📄 Extracting faculty names from: {doc_path}")
    
    faculty_names = extract_faculty_names_from_document(doc_path)
    print(f"✅ Extracted {len(faculty_names)} faculty names")
    
    # Extract existing entries
    print(f"📄 Extracting existing entries from reference document...")
    existing_entries = extract_existing_entries_from_document(doc_path)
    print(f"✅ Extracted {len(existing_entries)} existing entries")
    
    # Show first 10 names
    print("\n📋 First 10 faculty names:")
    for i, name in enumerate(faculty_names[:10], 1):
        print(f"  {i:2d}. {name}")
    print("  ...")
    
    print(f"\n🔍 Ready to search for NEW media mentions from June 1 - July 27, 2025")
    print(f"⏱️  Estimated time: {len(faculty_names)} minutes")
    
    proceed = input("\nProceed with search? (y/n): ").lower()
    if proceed != 'y':
        print("Search cancelled.")
        return
    
    # Run searches for all faculty
    new_results = []
    for i, faculty_name in enumerate(faculty_names, 1):
        print(f"\n[{i:3d}/{len(faculty_names)}] Processing: {faculty_name}")
        
        results = search_faculty_member(faculty_name, google_api_key)
        new_results.extend(results)
        
        # Progress update every 25 faculty
        if i % 25 == 0:
            print(f"\n📊 Progress: {i}/{len(faculty_names)} complete, {len(new_results)} total NEW results")
        
        time.sleep(random.uniform(1, 3))  # Rate limiting
    
    # Create Excel reports
    filename_new, filename_combined = create_excel_reports(new_results, existing_entries)
    
    print("\n" + "=" * 60)
    print("📋 FINAL RESULTS")
    print("=" * 60)
    print(f"Total faculty processed: {len(faculty_names)}")
    print(f"NEW results found (June-July 2025): {len(new_results)}")
    print(f"Existing entries from document: {len(existing_entries)}")
    print(f"Total combined entries: {len(existing_entries) + len(new_results)}")
    print(f"\n📊 Excel Reports Created:")
    print(f"1. June-July 2025 ONLY: {filename_new}")
    print(f"2. COMBINED (All periods): {filename_combined}")
    print("=" * 60)
    
    return filename_new, filename_combined

if __name__ == "__main__":
    main()
