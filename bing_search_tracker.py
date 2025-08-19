#!/usr/bin/env python3
"""
CSRR Faculty Media Tracker using Bing Search API
"""

import os
import requests
import json
import urllib.parse
import time
import random
import yaml
import pandas as pd
from docx import Document
from docx.shared import Inches
from datetime import datetime
from typing import List, Dict
import argparse

class BingFacultyMediaTracker:
    def __init__(self, config_file: str = "config.yaml"):
        self.config = self.load_config(config_file)
        self.bing_api_key = os.getenv("BING_API_KEY")
        
        if not self.bing_api_key:
            print("⚠️  Bing API not configured. Please set BING_API_KEY environment variable")
            print("💡 Get a free Bing Search API key from: https://www.microsoft.com/en-us/bing/apis/bing-web-search-api")
        
    def load_config(self, config_file: str) -> Dict:
        """Load configuration from YAML file"""
        default_config = {
            'search_period': {
                'start_date': '2025-06-01',
                'end_date': '2025-07-31'
            },
            'output': {
                'excel_filename': 'CSRR_Faculty_Media_Report.xlsx',
                'word_filename': 'CSRR_Faculty_Op-Eds.docx',
                'include_snippets': True,
                'max_results_per_faculty': 15,
                'save_to_downloads': True
            },
            'search': {
                'max_results_per_query': 10,
                'delay_between_searches': 1,
                'use_bing_api': True,
                'use_basic_search': True,
                'search_types': ['op-ed', 'interview', 'commentary', 'podcast', 'video']
            },
            'faculty': {
                'auto_fetch_from_website': True,
                'manual_list': []
            }
        }
        
        if os.path.exists(config_file):
            with open(config_file, 'r') as f:
                user_config = yaml.safe_load(f) or {}
                for key, value in user_config.items():
                    if key in default_config:
                        if isinstance(value, dict):
                            default_config[key].update(value)
                        else:
                            default_config[key] = value
                    else:
                        default_config[key] = value
        
        return default_config
    
    def search_bing_api(self, query: str, faculty_name: str) -> List[Dict]:
        """Search using Bing Web Search API"""
        if not self.bing_api_key:
            return []
        
        # Add date range to query
        start_date = self.config['search_period']['start_date']
        end_date = self.config['search_period']['end_date']
        date_query = f"{query} after:{start_date} before:{end_date}"
        
        url = "https://api.bing.microsoft.com/v7.0/search"
        headers = {
            "Ocp-Apim-Subscription-Key": self.bing_api_key
        }
        params = {
            "q": date_query,
            "count": self.config['search']['max_results_per_query'],
            "mkt": "en-US",
            "responseFilter": "Webpages"
        }
        
        try:
            response = requests.get(url, headers=headers, params=params, timeout=15)
            
            if response.status_code == 200:
                data = response.json()
                results = []
                
                if 'webPages' in data and 'value' in data['webPages']:
                    for item in data['webPages']['value']:
                        # Validate that faculty is mentioned
                        if self.validate_faculty_mention(faculty_name, item.get('name', ''), item.get('snippet', '')):
                            results.append({
                                'title': item.get('name', ''),
                                'url': item.get('url', ''),
                                'snippet': item.get('snippet', ''),
                                'source': self.extract_source(item.get('url', '')),
                                'date': self.extract_date(item.get('snippet', ''), item.get('url', '')),
                                'search_method': 'Bing API'
                            })
                
                return results
            else:
                print(f"  ⚠️  Bing API error: {response.status_code}")
                return []
                
        except Exception as e:
            print(f"  ⚠️  Bing API search error: {e}")
            return []
    
    def search_basic_web(self, query: str, faculty_name: str) -> List[Dict]:
        """Basic web scraping as fallback"""
        # This is a simplified version - you could implement more sophisticated scraping
        # For now, we'll return empty results to avoid rate limiting
        return []
    
    def search_faculty_media(self, faculty_name: str) -> List[Dict]:
        """Search for media mentions of a faculty member"""
        print(f"🔍 Searching for: {faculty_name}")
        
        # Build search queries
        search_queries = []
        for search_type in self.config['search']['search_types']:
            if search_type == 'op-ed':
                search_queries.extend([
                    f'"{faculty_name}" op-ed',
                    f'"{faculty_name}" opinion piece',
                    f'"{faculty_name}" editorial'
                ])
            elif search_type == 'interview':
                search_queries.extend([
                    f'"{faculty_name}" interview',
                    f'"{faculty_name}" interviewed'
                ])
            elif search_type == 'commentary':
                search_queries.extend([
                    f'"{faculty_name}" commentary',
                    f'"{faculty_name}" analysis',
                    f'"{faculty_name}" writes'
                ])
            elif search_type == 'podcast':
                search_queries.extend([
                    f'"{faculty_name}" podcast',
                    f'"{faculty_name}" audio interview'
                ])
            elif search_type == 'video':
                search_queries.extend([
                    f'"{faculty_name}" video',
                    f'"{faculty_name}" YouTube',
                    f'"{faculty_name}" television'
                ])
        
        all_results = []
        
        # Use Bing API if available
        if self.config['search']['use_bing_api'] and self.bing_api_key:
            for query in search_queries[:3]:  # Limit to top 3 queries
                try:
                    results = self.search_bing_api(query, faculty_name)
                    all_results.extend(results)
                    time.sleep(random.uniform(2, 3))  # Rate limiting
                except Exception as e:
                    print(f"  ⚠️  Bing API query error: {e}")
                    continue
        
        # Use basic search as fallback
        if self.config['search']['use_basic_search']:
            for query in search_queries[:2]:  # Limit basic search queries
                try:
                    results = self.search_basic_web(query, faculty_name)
                    all_results.extend(results)
                    time.sleep(random.uniform(1, self.config['search']['delay_between_searches']))
                except Exception as e:
                    print(f"  ⚠️  Basic search error: {e}")
                    continue
        
        # Remove duplicates and limit results
        unique_results = []
        seen_urls = set()
        
        for result in all_results:
            url = result.get('url', '')
            if url and url not in seen_urls:
                seen_urls.add(url)
                unique_results.append(result)
                
                if len(unique_results) >= self.config['output']['max_results_per_faculty']:
                    break
        
        if unique_results:
            print(f"  ✅ Found {len(unique_results)} articles")
        else:
            print(f"  ❌ No articles found")
        
        return unique_results
    
    def validate_faculty_mention(self, faculty_name: str, title: str, snippet: str) -> bool:
        """Validate that the faculty member is actually mentioned"""
        content_text = f"{title} {snippet}".lower()
        faculty_lower = faculty_name.lower()
        
        # Check for exact name match
        if faculty_lower in content_text:
            return True
        
        # Check for name components
        name_parts = faculty_lower.split()
        if len(name_parts) >= 2:
            last_name = name_parts[-1]
            first_name = name_parts[0]
            
            if last_name in content_text and first_name in content_text:
                title_words = title.lower().split()
                snippet_words = snippet.lower().split()
                all_words = title_words + snippet_words
                
                try:
                    first_pos = all_words.index(first_name)
                    last_pos = all_words.index(last_name)
                    if abs(first_pos - last_pos) <= 5:
                        return True
                except ValueError:
                    pass
        
        return False
    
    def extract_source(self, url: str) -> str:
        """Extract source name from URL"""
        try:
            domain = urllib.parse.urlparse(url).netloc
            source = domain.replace('www.', '').replace('.com', '').replace('.org', '').replace('.net', '').title()
            return source
        except:
            return "Unknown"
    
    def extract_date(self, content: str, url: str) -> str:
        """Extract publication date"""
        # Simple date extraction - could be enhanced
        try:
            # Look for common date patterns
            import re
            date_patterns = [
                r'\b\d{1,2}/\d{1,2}/\d{4}\b',
                r'\b\d{1,2}-\d{1,2}-\d{4}\b',
                r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{1,2},?\s+\d{4}\b'
            ]
            
            for pattern in date_patterns:
                match = re.search(pattern, content)
                if match:
                    return match.group()
            
            return "Date not found"
        except:
            return "Date not found"
    
    def fetch_faculty_list(self) -> List[str]:
        """Fetch faculty list from CSRR website"""
        if self.config['faculty']['auto_fetch_from_website']:
            try:
                print("🌐 Fetching faculty list from CSRR website...")
                faculty_list = self.scrape_faculty_from_website()
                if faculty_list:
                    print(f"✅ Found {len(faculty_list)} faculty members from website")
                    return faculty_list
            except Exception as e:
                print(f"⚠️  Could not fetch from website: {e}")
        
        # Fallback to comprehensive list
        fallback_list = [
            "Adil Haque", "Adnan Zulfiqar", "Alexander A. Reinert", "Alexander Hinton",
            "Alexis Karteron", "Ali A. Olomi", "Ali R. Chaudhary", "Ameena Ghaffar-Kucher",
            "Amir Hussain", "Anju Gupta", "Atalia Omer", "Atiya Aftab", "Audrey Truschke",
            "Aziz Rana", "Aziza Ahmed", "Behrooz Ghamari-Tabrizi", "Beth Baron", "Beth Stephens",
            "Bidisha Biswas", "Brittany Friedman", "Catherine M. Grosso", "Chaumtoli Huq",
            "Clark B. Lombardi", "Cyra A. Choudhury", "D. Asher Ghertner", "Dalia Fahmy",
            "Deepa Kumar", "Eid Mohamed", "Elise Boddie", "Ellen C. Yaroshefsky",
            "Emily Berman", "Emmaia Gelman", "Eric McDaniel", "Esaa Mohammad Samarah Samarsky",
            "Esther Canty-Barnes", "Faisal Kutty", "Faiza Sayed", "Falguni A. Sheth",
            "Farid Hafez", "Fatemeh Shams", "Gaiutra Devi Bahadur", "Ghada Ageel",
            "Hadi Khoshneviss", "Haider Ala Hamoudi", "Hajar Yazdiha", "Hatem Bazian",
            "Heba M. Khalil", "Huda J. Fakhreddine", "Irus Braverman", "Ivan Kalmar",
            "James R. Jones", "Jasbir K. Puar", "Jasmin Zine", "Jason Brownlee",
            "Jesse Norris", "Joel Beinin", "John Tehranian", "Jon Dubin", "Jonathan Feingold",
            "Jonathan Hafetz", "Jorge Contesse", "Juan Cole", "Karam Dana", "Karim Malak",
            "Karishma Desai", "Khaled A. Beydoun", "Khyati Joshi", "LaToya Baldwin Clark",
            "Lara Sheehi", "Leila Kawar", "Leyla Amzi-Erdogdular", "Lorenzo Veracini",
            "Mahruq Khan", "Marc O. Jones", "Margaret Hu", "Mark Bray", "Marta Esquilin",
            "Maryam Jamshidi", "Matthew Abraham", "Maura Finkelstein", "Mayte Green-Mercado",
            "Meera E. Deo", "Mitra Rastegar", "Mohammad Fadel", "Mojtaba Mahdavi",
            "Muhannad Ayyash", "Nader Hashemi", "Nadia Ahmad", "Nahed Samour",
            "Nancy A. Khalil", "Natsu Taylor Saito", "Nausheen Husain", "Naved Bakali",
            "Nazia Kazi", "Nermin Allam", "Norrinda Brown Hayat", "Noura Erakat",
            "Omar Al-Dewachi", "Omar Dajani", "Omar S. Dahi", "Omid Safi", "Ousseina Alidou",
            "Rabea Benhalim", "Rachel Godsil", "Raquel E Aldana", "Raz Segal",
            "Rebecca Hankins", "Riaz Tejani", "Robert S. Chang", "Sabreena Ghaffar-Siddiqui",
            "Sabrina Alimahomed", "Saeed Khan", "Sahar Mohamed Khamis", "Salman Sayyid",
            "Santiago Slabodsky", "Sarah Eltantawi", "Seema Saifee", "Sherene Razack",
            "Shirin Saeidi", "Shirin Sinnar", "Shoba Sivaprasad Wadhia", "Siham Elkassem",
            "Stacy Hawkins", "Stephen Dycus", "Stephen Sheehi", "Susan M. Akram",
            "Sylvia Chan-Malik", "Tahseen Shah", "Taleed El-Sabawi", "Tazeen M. Ali",
            "Timothy Eatman", "Timothy P. Daniels", "Timothy Raphael", "Toby Jones",
            "Udi Ofer", "Umayyah Cable", "Veena Dubal", "Victoria Ramenzoni",
            "Wadie Said", "Wendell Marsh", "Wendy Greene", "Whitney Strub",
            "William C. Banks", "William I. Robinson", "Yasmiyn Irizarry", "Zahra Ali",
            "Zain Abdullah", "Zakia Salime", "M. Shahid Alam", "Khalil Al-Anani",
            "Joseph Massad", "John L. Esposito", "Katherine M. Franke", "Tanya K. Hernandez"
        ]
        print(f"📋 Using fallback faculty list: {len(fallback_list)} members")
        return fallback_list
    
    def scrape_faculty_from_website(self) -> List[str]:
        """Scrape faculty list from CSRR website"""
        try:
            url = "https://csrr.rutgers.edu/about/faculty-affiliates/"
            response = requests.get(url, timeout=10)
            response.raise_for_status()
            
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Look for faculty names in various HTML elements
            faculty_names = []
            
            # Common patterns for faculty names
            name_selectors = [
                'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
                '.faculty-name', '.member-name', '.name',
                'strong', 'b', '.title'
            ]
            
            for selector in name_selectors:
                elements = soup.select(selector)
                for element in elements:
                    text = element.get_text().strip()
                    if text and len(text.split()) >= 2 and len(text) < 100:
                        # Basic validation for faculty names
                        if any(word in text.lower() for word in ['professor', 'dr', 'phd', 'esq']):
                            continue  # Skip titles
                        faculty_names.append(text)
            
            # Remove duplicates and clean up
            unique_names = list(set(faculty_names))
            return unique_names
            
        except Exception as e:
            print(f"Error scraping faculty from website: {e}")
            return []
    
    def create_excel_report(self, results: List[Dict]) -> str:
        """Create Excel report"""
        if not results:
            print("📊 No results to report")
            return ""
        
        df = pd.DataFrame(results)
        
        # Reorder columns
        columns = ['faculty_name', 'title', 'source', 'date', 'url', 'search_method']
        df = df.reindex(columns=columns, axis=1)
        
        filename = self.config['output']['excel_filename']
        df.to_excel(filename, index=False)
        print(f"📊 Excel report saved: {filename}")
        return filename
    
    def create_word_report(self, results: List[Dict]) -> str:
        """Create Word report"""
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "CSRR Faculty Media Coverage"
        doc.core_properties.author = "CSRR Media Tracker"
        
        # Add title
        doc.add_heading("Op-Eds by CSRR Faculty Affiliates", 0)
        
        # Add period
        start_date = self.config['search_period']['start_date']
        end_date = self.config['search_period']['end_date']
        period = f"{start_date} to {end_date}"
        doc.add_paragraph(period)
        
        # Group results by faculty
        faculty_results = {}
        for result in results:
            faculty = result.get('faculty_name', 'Unknown')
            if faculty not in faculty_results:
                faculty_results[faculty] = []
            faculty_results[faculty].append(result)
        
        # Add faculty sections
        for faculty_name in sorted(faculty_results.keys()):
            doc.add_paragraph(faculty_name)
            
            for result in faculty_results[faculty_name]:
                title = result.get('title', 'No title')
                source = result.get('source', 'Unknown source')
                date = result.get('date', 'No date')
                url = result.get('url', '')
                
                entry = f"{faculty_name}, {title}, {source}, {date}, {url}."
                doc.add_paragraph(entry)
        
        # Save to Downloads folder if configured
        if self.config['output']['save_to_downloads']:
            downloads_path = os.path.expanduser("~/Downloads")
            filename = os.path.join(downloads_path, self.config['output']['word_filename'])
        else:
            filename = self.config['output']['word_filename']
        
        doc.save(filename)
        print(f"📄 Word report saved: {filename}")
        return filename
    
    def run_search(self) -> List[Dict]:
        """Run the complete search process"""
        print("=" * 60)
        print("🎯 BING CSRR FACULTY MEDIA TRACKER")
        print("=" * 60)
        
        if self.bing_api_key:
            print("✅ Bing Search API configured - Enhanced search enabled")
        else:
            print("⚠️  Bing API not configured. Using basic search only.")
            print("💡 To enable enhanced search, set BING_API_KEY environment variable")
        
        # Get faculty list
        faculty_list = self.fetch_faculty_list()
        
        if self.config['faculty']['manual_list']:
            faculty_list = self.config['faculty']['manual_list']
        
        print(f"👥 Processing {len(faculty_list)} faculty members")
        print(f"📅 Period: {self.config['search_period']['start_date']} to {self.config['search_period']['end_date']}")
        print("=" * 60)
        
        all_results = []
        faculty_with_results = 0
        
        for i, faculty_name in enumerate(faculty_list, 1):
            print(f"[{i:3d}/{len(faculty_list)}] ", end="")
            
            results = self.search_faculty_media(faculty_name)
            
            # Add faculty name to results
            for result in results:
                result['faculty_name'] = faculty_name
            
            all_results.extend(results)
            
            if results:
                faculty_with_results += 1
            
            # Progress update every 20 faculty
            if i % 20 == 0:
                print(f"\n📊 Progress: {i}/{len(faculty_list)} faculty processed")
                print(f"   Found articles for: {faculty_with_results} faculty")
                print(f"   Total articles: {len(all_results)}")
                print()
        
        print("=" * 60)
        print("📋 BING SEARCH COMPLETED")
        print("=" * 60)
        print(f"Total faculty processed: {len(faculty_list)}")
        print(f"Faculty with articles: {faculty_with_results}")
        print(f"Total articles found: {len(all_results)}")
        print(f"Success rate: {(faculty_with_results/len(faculty_list)*100):.1f}%")
        
        return all_results

def main():
    parser = argparse.ArgumentParser(description="CSRR Faculty Media Tracker using Bing Search")
    parser.add_argument("--config", default="config.yaml", help="Configuration file path")
    parser.add_argument("--quick-test", action="store_true", help="Run quick test with first 5 faculty")
    parser.add_argument("--setup-api", action="store_true", help="Show API setup instructions")
    
    args = parser.parse_args()
    
    if args.setup_api:
        print("🔧 BING SEARCH API SETUP INSTRUCTIONS")
        print("=" * 50)
        print("1. Go to: https://www.microsoft.com/en-us/bing/apis/bing-web-search-api")
        print("2. Click 'Try now' or 'Get started'")
        print("3. Sign in with your Microsoft account")
        print("4. Create a new resource (choose 'Web Search API')")
        print("5. Get your API key from the 'Keys and Endpoint' section")
        print("6. Set the environment variable:")
        print("   export BING_API_KEY='your_api_key_here'")
        print("7. Free tier: 1,000 searches/month")
        print("8. Paid tier: $3 per 1,000 searches")
        return
    
    tracker = BingFacultyMediaTracker(args.config)
    
    if args.quick_test:
        print("🧪 Running quick test with first 5 faculty members...")
        faculty_list = tracker.fetch_faculty_list()
        test_faculty = faculty_list[:5]
        tracker.config['faculty']['manual_list'] = test_faculty
        tracker.config['faculty']['auto_fetch_from_website'] = False
    
    results = tracker.run_search()
    
    if results:
        # Create reports
        excel_file = tracker.create_excel_report(results)
        word_file = tracker.create_word_report(results)
        
        print(f"\n✅ REPORTS GENERATED:")
        print(f"📊 Excel: {excel_file}")
        print(f"📄 Word: {word_file}")
    else:
        print("\n❌ No results found. Consider:")
        print("   - Checking your API key")
        print("   - Adjusting search dates")
        print("   - Using different search terms")
    
    print("\n🎉 Bing search completed successfully!")
    print("📁 Check the generated files for your results.")

if __name__ == "__main__":
    main()
