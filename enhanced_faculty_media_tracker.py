#!/usr/bin/env python3
"""
CSSR_FacultyOpEds_AutomationTool - Enhanced Faculty Media Tracker
Comprehensive search for op-eds, interviews, and media appearances

Author: AI Assistant
Date: 2025
"""

import pandas as pd
from docx import Document
import time
import random
from datetime import datetime, timedelta
import requests
from bs4 import BeautifulSoup
import urllib.parse
import re
import json
import os
from typing import List, Dict, Optional
import argparse
import yaml
from pathlib import Path

class EnhancedFacultyMediaTracker:
    """Enhanced automated tool for tracking CSRR faculty media appearances"""
    
    def __init__(self, config_file: str = "config.yaml"):
        self.config = self.load_config(config_file)
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        })
        self.results = []
        
        # API Configuration
        self.google_api_key = os.getenv("GOOGLE_API_KEY")
        self.google_cse_id = os.getenv("GOOGLE_CSE_ID")
        
        if not self.google_api_key or not self.google_cse_id:
            print("⚠️  Google API not configured. Using basic search only.")
            print("💡 To enable enhanced search, set GOOGLE_API_KEY and GOOGLE_CSE_ID environment variables")
        
    def load_config(self, config_file: str) -> Dict:
        """Load configuration from YAML file"""
        default_config = {
            'search_period': {
                'start_date': '2025-06-01',
                'end_date': '2025-07-31'
            },
            'output': {
                'excel_filename': 'CSRR_Faculty_Media_Report.xlsx',
                'word_filename': 'CSRR_Faculty_Op-Eds.docx',
                'include_snippets': True,
                'max_results_per_faculty': 10,  # Increased for enhanced search
                'save_to_downloads': True  # New option to save to Downloads folder
            },
            'search': {
                'max_results_per_query': 10,  # Increased for Google API
                'delay_between_searches': 1,  # Reduced for API
                'use_google_api': True,
                'use_basic_search': True,  # Fallback
                'search_types': ['op-ed', 'interview', 'commentary', 'podcast', 'video']
            },
            'faculty': {
                'auto_fetch_from_website': True,
                'manual_list': []
            }
        }
        
        if os.path.exists(config_file):
            with open(config_file, 'r') as f:
                user_config = yaml.safe_load(f) or {}
                for key, value in user_config.items():
                    if key in default_config:
                        if isinstance(value, dict):
                            default_config[key].update(value)
                        else:
                            default_config[key] = value
                    else:
                        default_config[key] = value
        
        return default_config
    
    def create_default_config(self, config_file: str = "config.yaml") -> None:
        """Create a default configuration file"""
        default_config = {
            'search_period': {
                'description': 'Date range for media search (YYYY-MM-DD format)',
                'start_date': '2025-06-01',
                'end_date': '2025-08-19'
            },
            'output': {
                'description': 'Output file settings',
                'excel_filename': 'CSRR_Faculty_Media_Report.xlsx',
                'word_filename': 'CSRR_Faculty_Op-Eds_May31_to_Aug19_2025.docx',
                'include_snippets': True,
                'max_results_per_faculty': 15,
                'save_to_downloads': True
            },
            'search': {
                'description': 'Search behavior settings',
                'max_results_per_query': 10,
                'delay_between_searches': 5,
                'use_google_api': True,
                'use_basic_search': False,
                'trusted_sources_only': False,
                'search_types': ['op-ed', 'interview', 'commentary', 'podcast', 'video']
            },
            'faculty': {
                'description': 'Faculty list settings',
                'auto_fetch_from_website': True,
                'manual_list': []
            }
        }
        
        with open(config_file, 'w') as f:
            yaml.dump(default_config, f, default_flow_style=False, indent=2)
        
        print(f"✅ Default configuration created: {config_file}")
        print("💡 Edit this file to customize search parameters")
    
    def fetch_faculty_list(self) -> List[str]:
        """Fetch faculty list from CSRR website"""
        if self.config['faculty']['auto_fetch_from_website']:
            try:
                print("🌐 Fetching faculty list from CSRR website...")
                faculty_list = self.scrape_faculty_from_website()
                if faculty_list:
                    print(f"✅ Found {len(faculty_list)} faculty members from website")
                    return faculty_list
            except Exception as e:
                print(f"⚠️  Could not fetch from website: {e}")
        
        # Fallback to comprehensive list
        fallback_list = [
            "Adil Haque", "Adnan Zulfiqar", "Alexander A. Reinert", "Alexander Hinton",
            "Alexis Karteron", "Ali A. Olomi", "Ali R. Chaudhary", "Ameena Ghaffar-Kucher",
            "Amir Hussain", "Anju Gupta", "Atalia Omer", "Atiya Aftab", "Audrey Truschke",
            "Aziz Rana", "Aziza Ahmed", "Behrooz Ghamari-Tabrizi", "Beth Baron", "Beth Stephens",
            "Bidisha Biswas", "Brittany Friedman", "Catherine M. Grosso", "Chaumtoli Huq",
            "Clark B. Lombardi", "Cyra A. Choudhury", "D. Asher Ghertner", "Dalia Fahmy",
            "Deepa Kumar", "Eid Mohamed", "Elise Boddie", "Ellen C. Yaroshefsky",
            "Emily Berman", "Emmaia Gelman", "Eric McDaniel", "Esaa Mohammad Samarah Samarsky",
            "Esther Canty-Barnes", "Faisal Kutty", "Faiza Sayed", "Falguni A. Sheth",
            "Farid Hafez", "Fatemeh Shams", "Gaiutra Devi Bahadur", "Ghada Ageel",
            "Hadi Khoshneviss", "Haider Ala Hamoudi", "Hajar Yazdiha", "Hatem Bazian",
            "Heba M. Khalil", "Huda J. Fakhreddine", "Irus Braverman", "Ivan Kalmar",
            "James R. Jones", "Jasbir K. Puar", "Jasmin Zine", "Jason Brownlee",
            "Jesse Norris", "Joel Beinin", "John Tehranian", "Jon Dubin", "Jonathan Feingold",
            "Jonathan Hafetz", "Jorge Contesse", "Juan Cole", "Karam Dana", "Karim Malak",
            "Karishma Desai", "Khaled A. Beydoun", "Khyati Joshi", "LaToya Baldwin Clark",
            "Lara Sheehi", "Leila Kawar", "Leyla Amzi-Erdogdular", "Lorenzo Veracini",
            "Mahruq Khan", "Marc O. Jones", "Margaret Hu", "Mark Bray", "Marta Esquilin",
            "Maryam Jamshidi", "Matthew Abraham", "Maura Finkelstein", "Mayte Green-Mercado",
            "Meera E. Deo", "Mitra Rastegar", "Mohammad Fadel", "Mojtaba Mahdavi",
            "Muhannad Ayyash", "Nader Hashemi", "Nadia Ahmad", "Nahed Samour",
            "Nancy A. Khalil", "Natsu Taylor Saito", "Nausheen Husain", "Naved Bakali",
            "Nazia Kazi", "Nermin Allam", "Norrinda Brown Hayat", "Noura Erakat",
            "Omar Al-Dewachi", "Omar Dajani", "Omar S. Dahi", "Omid Safi", "Ousseina Alidou",
            "Rabea Benhalim", "Rachel Godsil", "Raquel E Aldana", "Raz Segal",
            "Rebecca Hankins", "Riaz Tejani", "Robert S. Chang", "Sabreena Ghaffar-Siddiqui",
            "Sabrina Alimahomed", "Saeed Khan", "Sahar Mohamed Khamis", "Salman Sayyid",
            "Santiago Slabodsky", "Sarah Eltantawi", "Seema Saifee", "Sherene Razack",
            "Shirin Saeidi", "Shirin Sinnar", "Shoba Sivaprasad Wadhia", "Siham Elkassem",
            "Stacy Hawkins", "Stephen Dycus", "Stephen Sheehi", "Susan M. Akram",
            "Sylvia Chan-Malik", "Tahseen Shah", "Taleed El-Sabawi", "Tazeen M. Ali",
            "Timothy Eatman", "Timothy P. Daniels", "Timothy Raphael", "Toby Jones",
            "Udi Ofer", "Umayyah Cable", "Veena Dubal", "Victoria Ramenzoni",
            "Wadie Said", "Wendell Marsh", "Wendy Greene", "Whitney Strub",
            "William C. Banks", "William I. Robinson", "Yasmiyn Irizarry", "Zahra Ali",
            "Zain Abdullah", "Zakia Salime", "M. Shahid Alam", "Khalil Al-Anani",
            "Joseph Massad", "John L. Esposito", "Katherine M. Franke", "Tanya K. Hernandez"
        ]
        print(f"📋 Using fallback faculty list: {len(fallback_list)} members")
        return fallback_list
    
    def scrape_faculty_from_website(self) -> List[str]:
        """Scrape faculty names from CSRR website"""
        url = 'https://csrr.rutgers.edu/about/faculty-affiliates/'
        response = self.session.get(url)
        soup = BeautifulSoup(response.content, 'html.parser')
        
        faculty_names = []
        text_content = soup.get_text()
        lines = text_content.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            next_line = lines[i+1].strip() if i+1 < len(lines) else ""
            
            if any(title in next_line.lower() for title in [
                'professor', 'associate professor', 'assistant professor', 
                'clinical professor', 'emeritus professor', 'visiting professor',
                'distinguished professor', 'lecturer', 'instructor'
            ]):
                name = line.strip()
                name = re.sub(r'^(Dr\.|Prof\.|Professor)\s+', '', name)
                name = re.sub(r'\s+(Jr\.|Sr\.|II|III|IV)$', r' \1', name)
                
                if len(name.split()) >= 2 and len(name) < 50:
                    faculty_names.append(name)
        
        for element in soup.find_all(['h3', 'h4', 'h5', 'strong', 'b']):
            text = element.get_text().strip()
            if text and len(text.split()) >= 2 and len(text) < 50:
                words = text.split()
                if all(word[0].isupper() for word in words if word):
                    faculty_names.append(text)
        
        faculty_names = sorted(list(set(faculty_names)))
        filtered_faculty = []
        
        for name in faculty_names:
            if not any(term in name.upper() for term in ['NEWS', 'UPDATE', 'ALERT', 'CLICK', 'VIEW', 'READ']):
                name_parts = name.split()
                if len(name_parts) >= 2 and all(part[0].isupper() and part[1:].islower() for part in name_parts if part.isalpha()):
                    filtered_faculty.append(name)
        
        return filtered_faculty
    
    def search_google_api(self, query: str, faculty_name: str) -> List[Dict]:
        """Search using Google Custom Search API"""
        if not self.google_api_key or not self.google_cse_id:
            return []
        
        try:
            # Add date range to query
            start_date = self.config['search_period']['start_date']
            end_date = self.config['search_period']['end_date']
            date_query = f"{query} after:{start_date} before:{end_date}"
            
            url = "https://www.googleapis.com/customsearch/v1"
            params = {
                'key': self.google_api_key,
                'cx': self.google_cse_id,
                'q': date_query,
                'num': min(self.config['search']['max_results_per_query'], 10),  # Google API max is 10
                'dateRestrict': 'm1'  # Restrict to last month
            }
            
            response = self.session.get(url, params=params, timeout=15)
            response.raise_for_status()
            
            data = response.json()
            results = []
            
            if 'items' in data:
                for item in data['items']:
                    title = item.get('title', '')
                    link = item.get('link', '')
                    snippet = item.get('snippet', '')
                    
                    # Validate faculty mention
                    if not self.validate_faculty_mention(faculty_name, title, snippet):
                        continue
                    
                    # Filter out irrelevant sources
                    if not self.is_relevant_source(link, title, snippet):
                        continue
                    
                    # Extract source and date
                    source = self.extract_source(link)
                    pub_date = self.extract_date(f"{title} {snippet}", link)
                    
                    # Filter out articles with unknown dates
                    if pub_date == "Unknown":
                        continue
                    
                    results.append({
                        'faculty_name': faculty_name,
                        'title': title,
                        'url': link,
                        'snippet': snippet,
                        'source': source,
                        'publication_date': pub_date,
                        'search_method': 'Google API'
                    })
            
            return results
            
        except Exception as e:
            print(f"  ⚠️  Google API search error: {e}")
            return []
    
    def search_basic_web(self, query: str, faculty_name: str) -> List[Dict]:
        """Basic web search using Bing (fallback)"""
        try:
            start_date = self.config['search_period']['start_date']
            end_date = self.config['search_period']['end_date']
            date_query = f"{query} after:{start_date} before:{end_date}"
            
            url = f"https://www.bing.com/search?q={urllib.parse.quote(date_query)}&count={self.config['search']['max_results_per_query']}"
            
            response = self.session.get(url, timeout=15)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            results = []
            
            for result in soup.find_all('li', class_='b_algo')[:self.config['search']['max_results_per_query']]:
                title_elem = result.find('h2')
                if not title_elem:
                    continue
                    
                title_link = title_elem.find('a')
                if not title_link:
                    continue
                
                title = title_link.get_text(strip=True)
                link = title_link.get('href', '')
                
                snippet_elem = result.find('p')
                snippet = snippet_elem.get_text(strip=True) if snippet_elem else ''
                
                if not self.validate_faculty_mention(faculty_name, title, snippet):
                    continue
                
                # Filter out irrelevant sources
                if not self.is_relevant_source(link, title, snippet):
                    continue
                
                source = self.extract_source(link)
                pub_date = self.extract_date(f"{title} {snippet}", link)
                
                # Filter out articles with unknown dates
                if pub_date == "Unknown":
                    continue
                
                results.append({
                    'faculty_name': faculty_name,
                    'title': title,
                    'url': link,
                    'snippet': snippet,
                    'source': source,
                    'publication_date': pub_date,
                    'search_method': 'Basic Web'
                })
            
            return results
            
        except Exception as e:
            print(f"  ⚠️  Basic web search error: {e}")
            return []
    
    def search_faculty_media(self, faculty_name: str) -> List[Dict]:
        """Comprehensive search for faculty media appearances"""
        print(f"🔍 Searching for: {faculty_name}")
        
        # Enhanced search queries focused on op-eds, print interviews, and television interviews
        search_queries = []
        
        # Op-eds and opinion pieces
        search_queries.extend([
            f'"{faculty_name}" op-ed',
            f'"{faculty_name}" opinion piece',
            f'"{faculty_name}" editorial',
            f'"{faculty_name}" opinion',
            f'"{faculty_name}" guest column',
            f'"{faculty_name}" commentary'
        ])
        
        # Print interviews
        search_queries.extend([
            f'"{faculty_name}" interview',
            f'"{faculty_name}" interviewed',
            f'"{faculty_name}" speaks with',
            f'"{faculty_name}" Q&A',
            f'"{faculty_name}" conversation with'
        ])
        
        # Television interviews
        search_queries.extend([
            f'"{faculty_name}" television',
            f'"{faculty_name}" TV interview',
            f'"{faculty_name}" news interview',
            f'"{faculty_name}" CNN',
            f'"{faculty_name}" MSNBC',
            f'"{faculty_name}" Fox News',
            f'"{faculty_name}" PBS',
            f'"{faculty_name}" ABC',
            f'"{faculty_name}" CBS',
            f'"{faculty_name}" NBC'
        ])
        
        all_results = []
        
        # Use Google API if available
        if self.config['search']['use_google_api'] and self.google_api_key:
            for query in search_queries[:3]:  # Limit to top 3 queries to avoid rate limits
                try:
                    results = self.search_google_api(query, faculty_name)
                    all_results.extend(results)
                    time.sleep(random.uniform(2, 3))  # Increased API rate limiting
                except Exception as e:
                    if "429" in str(e):
                        print(f"  ⚠️  Rate limit hit, waiting 10 seconds...")
                        time.sleep(10)  # Wait longer on rate limit
                    else:
                        print(f"  ⚠️  Google API query error: {e}")
                    continue
        
        # Use basic search as fallback or supplement
        if self.config['search']['use_basic_search']:
            for query in search_queries[:3]:  # Limit basic search queries
                try:
                    results = self.search_basic_web(query, faculty_name)
                    all_results.extend(results)
                    time.sleep(random.uniform(1, self.config['search']['delay_between_searches']))
                except Exception as e:
                    print(f"  ⚠️  Basic search error: {e}")
                    continue
        
        # Remove duplicates and limit results
        unique_results = []
        seen_urls = set()
        
        for result in all_results:
            url = result.get('url', '')
            if url and url not in seen_urls:
                seen_urls.add(url)
                unique_results.append(result)
                
                if len(unique_results) >= self.config['output']['max_results_per_faculty']:
                    break
        
        if unique_results:
            print(f"  ✅ Found {len(unique_results)} articles")
        else:
            print(f"  ❌ No articles found")
        
        return unique_results
    
    def is_relevant_source(self, url: str, title: str, snippet: str) -> bool:
        """Strict filtering for ONLY op-eds, print interviews, and television interviews"""
        url_lower = url.lower()
        title_lower = title.lower()
        snippet_lower = snippet.lower()
        content_lower = f"{title_lower} {snippet_lower}"
        
        # STRICT: Must be from legitimate news/media sources
        legitimate_domains = [
            'nytimes.com', 'washingtonpost.com', 'wsj.com', 'usatoday.com', 'latimes.com',
            'chicagotribune.com', 'bostonglobe.com', 'philly.com', 'miamiherald.com',
            'cnn.com', 'msnbc.com', 'foxnews.com', 'abcnews.go.com', 'cbsnews.com', 'nbcnews.com',
            'pbs.org', 'npr.org', 'bbc.com', 'reuters.com', 'ap.org', 'bloomberg.com',
            'politico.com', 'thehill.com', 'rollcall.com', 'nationalreview.com', 'newyorker.com',
            'atlantic.com', 'huffpost.com', 'vox.com', 'slate.com', 'salon.com',
            'guardian.com', 'independent.co.uk', 'telegraph.co.uk', 'ft.com', 'economist.com',
            'aljazeera.com', 'middleeasteye.net', 'newarab.com', 'arabnews.com',
            'law.com', 'abajournal.com', 'law360.com', 'scotusblog.com', 'justsecurity.org',
            'lawfaremedia.org', 'balkinization.net', 'volokh.com', 'concurringopinions.com'
        ]
        
        # Check if URL contains any legitimate domain
        is_legitimate_source = any(domain in url_lower for domain in legitimate_domains)
        if not is_legitimate_source:
            return False
        
        # STRICT: Must contain specific media content indicators
        media_indicators = [
            'op-ed', 'opinion', 'editorial', 'guest column', 'commentary',
            'interview', 'interviewed', 'speaks with', 'conversation with', 'q&a',
            'television', 'tv interview', 'news interview', 'appears on', 'discusses',
            'writes', 'author', 'byline', 'contributed', 'analysis'
        ]
        
        has_media_indicator = any(indicator in content_lower for indicator in media_indicators)
        if not has_media_indicator:
            return False
        
        # EXCLUDE: Social media, academic papers, irrelevant content
        exclude_patterns = [
            'facebook.com', 'instagram.com', 'tiktok.com', 'twitter.com', 'x.com',
            'linkedin.com', 'reddit.com', 'youtube.com', 'researchgate.net', 'jstor.org',
            'academia.edu', 'scholar.google.com', 'arxiv.org', 'ssrn.com',
            'archive', 'archives', 'course', 'syllabus', 'academic', 'student',
            'sale', 'shop', 'store', 'product', 'booking', 'hotel', 'travel',
            'obituary', 'funeral', 'memorial', 'wedding', 'birthday', 'party'
        ]
        
        for pattern in exclude_patterns:
            if pattern in url_lower or pattern in content_lower:
                return False
        
        return True
    
    def validate_faculty_mention(self, faculty_name: str, title: str, snippet: str) -> bool:
        """Validate that the faculty member is actually mentioned"""
        content_text = f"{title} {snippet}".lower()
        faculty_lower = faculty_name.lower()
        
        # Check for exact name match
        if faculty_lower in content_text:
            return True
        
        # Check for name components
        name_parts = faculty_lower.split()
        if len(name_parts) >= 2:
            last_name = name_parts[-1]
            first_name = name_parts[0]
            
            if last_name in content_text and first_name in content_text:
                title_words = title.lower().split()
                snippet_words = snippet.lower().split()
                all_words = title_words + snippet_words
                
                try:
                    first_pos = all_words.index(first_name)
                    last_pos = all_words.index(last_name)
                    if abs(first_pos - last_pos) <= 5:
                        return True
                except ValueError:
                    pass
        
        return False
    
    def extract_source(self, url: str) -> str:
        """Extract source name from URL"""
        try:
            domain = urllib.parse.urlparse(url).netloc
            source = domain.replace('www.', '').replace('.com', '').replace('.org', '').replace('.net', '').title()
            return source
        except:
            return "Unknown"
    
    def extract_date(self, content: str, url: str) -> str:
        """Extract publication date with strict validation for June 1 - Aug 19, 2025"""
        # Only look for dates in 2025, specifically June-August
        date_patterns = [
            r'(Jun|Jul|Aug)[a-z]*\.?\s+\d{1,2},?\s+2025',
            r'\d{1,2}[/-](06|07|08)[/-]2025',
            r'2025-(06|07|08)-\d{2}',
            r'(June|July|August)\s+\d{1,2},?\s+2025'
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                date_str = match.group(0)
                # Additional validation - must be June 1 or later, Aug 19 or earlier
                if self.is_valid_date_in_range(date_str):
                    return date_str
        
        # If no valid date found, return "Unknown" instead of generic "2025"
        return "Unknown"
    
    def is_valid_date_in_range(self, date_str: str) -> bool:
        """Validate that date is within June 1 - August 19, 2025"""
        try:
            # Parse the date string
            date_str_clean = date_str.strip()
            
            # Try different formats
            for fmt in ['%b %d, %Y', '%B %d, %Y', '%m/%d/%Y', '%Y-%m-%d']:
                try:
                    parsed_date = datetime.strptime(date_str_clean, fmt)
                    
                    # Check if it's in our range: June 1 - August 19, 2025
                    start_date = datetime(2025, 6, 1)
                    end_date = datetime(2025, 8, 19)
                    
                    if start_date <= parsed_date <= end_date:
                        return True
                    else:
                        return False
                        
                except ValueError:
                    continue
            
            return False
        except:
            return False
    
    def create_excel_report(self, results: List[Dict]) -> str:
        """Create Excel report"""
        if not results:
            df = pd.DataFrame(columns=[
                "Faculty Name", "Title", "Source", "URL", 
                "Publication Date", "Date Found", "Snippet", "Search Method"
            ])
        else:
            excel_data = []
            for result in results:
                excel_data.append({
                    "Faculty Name": result['faculty_name'],
                    "Title": result['title'],
                    "Source": result['source'],
                    "URL": result['url'],
                    "Publication Date": result['publication_date'],
                    "Date Found": datetime.now().strftime('%Y-%m-%d'),
                    "Snippet": result['snippet'][:500] + "..." if len(result['snippet']) > 500 else result['snippet'],
                    "Search Method": result.get('search_method', 'Unknown')
                })
            df = pd.DataFrame(excel_data)
        
        filename = self.config['output']['excel_filename']
        df.to_excel(filename, index=False)
        print(f"📊 Excel report saved: {filename}")
        return filename
    
    def create_word_report(self, results: List[Dict]) -> str:
        """Create Word report with clean, professional formatting"""
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "CSRR Faculty Op-Eds, Print Interviews, and Television Interviews"
        doc.core_properties.author = "CSRR Enhanced Media Tracker"
        
        # Main title
        title_para = doc.add_heading("CSRR Faculty Op-Eds, Print Interviews, and Television Interviews", 0)
        
        # Subtitle with date range
        start_date = self.config['search_period']['start_date']
        end_date = self.config['search_period']['end_date']
        
        from datetime import datetime
        try:
            start_dt = datetime.strptime(start_date, '%Y-%m-%d')
            end_dt = datetime.strptime(end_date, '%Y-%m-%d')
            # Use cross-platform date formatting
            period = f"{start_dt.strftime('%B %d, %Y')} - {end_dt.strftime('%B %d, %Y')}"
        except:
            period = f"{start_date} to {end_date}"
        
        subtitle_para = doc.add_paragraph(period)
        doc.add_paragraph("")
        
        # Group results by faculty
        faculty_results = {}
        for result in results:
            faculty = result['faculty_name']
            if faculty not in faculty_results:
                faculty_results[faculty] = []
            faculty_results[faculty].append(result)
        
        # Add results by faculty
        for faculty_name in sorted(faculty_results.keys()):
            # Add faculty name as a paragraph
            faculty_para = doc.add_paragraph(faculty_name)
            
            # Add articles for this faculty
            results = faculty_results[faculty_name]
            for result in results:
                # Clean up the title (remove extra spaces and truncate if too long)
                title = result['title'].strip()
                if len(title) > 100:
                    title = title[:97] + "..."
                
                # Format the entry cleanly
                formatted_entry = (
                    f"{title}, "
                    f"{result['source']}, "
                    f"{result['publication_date']}, "
                    f"{result['url']}."
                )
                doc.add_paragraph(formatted_entry)
            
            # Add space between faculty
            doc.add_paragraph("")
        
        # Add faculty with no results
        all_faculty = set(self.fetch_faculty_list())
        faculty_with_results = set(faculty_results.keys())
        faculty_without_results = all_faculty - faculty_with_results
        
        for faculty_name in sorted(faculty_without_results):
            doc.add_paragraph(faculty_name)
        
        # Create filename with date range
        start_date = self.config['search_period']['start_date']
        end_date = self.config['search_period']['end_date']
        
        from datetime import datetime
        try:
            start_dt = datetime.strptime(start_date, '%Y-%m-%d')
            end_dt = datetime.strptime(end_date, '%Y-%m-%d')
            
            if start_dt.year == end_dt.year and start_dt.month == 1 and start_dt.day == 1:
                date_suffix = f"{start_dt.year}"
            else:
                date_suffix = f"{start_dt.strftime('%b%Y')}_to_{end_dt.strftime('%b%Y')}"
        except:
            date_suffix = f"{start_date.replace('-', '')}_to_{end_date.replace('-', '')}"
        
        filename = f"CSRR_Faculty_Op-Eds_{date_suffix}.docx"
        
        # Save to Downloads folder if configured
        if self.config['output']['save_to_downloads']:
            downloads_path = os.path.expanduser("~/Downloads")
            filename = os.path.join(downloads_path, filename)
        
        doc.save(filename)
        print(f"📄 Word report saved: {filename}")
        return filename
    
    def run_search(self) -> Dict[str, str]:
        """Run the complete enhanced media search"""
        print("=" * 60)
        print("🎯 ENHANCED CSRR FACULTY MEDIA TRACKER")
        print("=" * 60)
        
        # Check API configuration
        if self.google_api_key and self.google_cse_id:
            print("✅ Google Custom Search API configured - Enhanced search enabled")
        else:
            print("⚠️  Google API not configured - Using basic search only")
            print("💡 For enhanced results, set GOOGLE_API_KEY and GOOGLE_CSE_ID")
        
        # Load faculty list
        faculty_list = self.fetch_faculty_list()
        print(f"👥 Processing {len(faculty_list)} faculty members")
        print(f"📅 Period: {self.config['search_period']['start_date']} to {self.config['search_period']['end_date']}")
        print("=" * 60)
        print()
        
        all_results = []
        faculty_with_results = 0
        
        # Process each faculty member
        for i, faculty_name in enumerate(faculty_list, 1):
            print(f"[{i:3d}/{len(faculty_list)}] ", end="")
            
            results = self.search_faculty_media(faculty_name)
            
            if results:
                faculty_with_results += 1
                all_results.extend(results)
            
            # Progress update
            if i % 20 == 0:
                print(f"\n📊 Progress: {i}/{len(faculty_list)} faculty processed")
                print(f"   Found articles for: {faculty_with_results} faculty")
                print(f"   Total articles: {len(all_results)}")
                print()
        
        # Final statistics
        print("\n" + "=" * 60)
        print("📋 ENHANCED SEARCH COMPLETED")
        print("=" * 60)
        print(f"Total faculty processed: {len(faculty_list)}")
        print(f"Faculty with articles: {faculty_with_results}")
        print(f"Total articles found: {len(all_results)}")
        print(f"Success rate: {faculty_with_results/len(faculty_list)*100:.1f}%")
        
        # Search method breakdown
        if all_results:
            search_methods = {}
            for result in all_results:
                method = result.get('search_method', 'Unknown')
                search_methods[method] = search_methods.get(method, 0) + 1
            
            print(f"\n🔍 Search Method Breakdown:")
            for method, count in search_methods.items():
                print(f"   {method}: {count} articles")
        
        # Generate reports
        excel_file = self.create_excel_report(all_results)
        word_file = self.create_word_report(all_results)
        
        print(f"\n✅ REPORTS GENERATED:")
        print(f"📊 Excel: {excel_file}")
        print(f"📄 Word: {word_file}")
        
        return {
            'excel': excel_file,
            'word': word_file,
            'total_articles': len(all_results),
            'faculty_with_results': faculty_with_results
        }

def main():
    """Main function with command line interface"""
    parser = argparse.ArgumentParser(description='CSSR_FacultyOpEds_AutomationTool - Enhanced CSRR Faculty Media Tracker')
    parser.add_argument('--config', default='config.yaml', help='Configuration file path')
    parser.add_argument('--create-config', action='store_true', help='Create a default configuration file')
    parser.add_argument('--quick-test', action='store_true', help='Run a quick test with first 5 faculty')
    parser.add_argument('--setup-api', action='store_true', help='Show API setup instructions')
    
    args = parser.parse_args()
    
    if args.setup_api:
        print("=" * 60)
        print("🔧 GOOGLE CUSTOM SEARCH API SETUP")
        print("=" * 60)
        print("\n📋 Steps to enable enhanced search:")
        print("\n1. Go to Google Cloud Console: https://console.cloud.google.com/")
        print("2. Create a new project or select existing one")
        print("3. Enable Custom Search API")
        print("4. Create credentials (API Key)")
        print("5. Go to Custom Search Engine: https://cse.google.com/")
        print("6. Create a new search engine")
        print("7. Set environment variables:")
        print("   export GOOGLE_API_KEY='your_api_key_here'")
        print("   export GOOGLE_CSE_ID='your_cse_id_here'")
        print("\n💰 Cost: ~$5-10 per month for thorough searching")
        print("📈 Coverage: Comprehensive web search including news, social media, podcasts")
        return
    
    if args.create_config:
        tracker = EnhancedFacultyMediaTracker()
        tracker.create_default_config(args.config)
        return
    
    # Initialize tracker
    tracker = EnhancedFacultyMediaTracker(args.config)
    
    if args.quick_test:
        print("🧪 Running quick test with first 5 faculty members...")
        # Get the actual faculty list from website
        faculty_list = tracker.fetch_faculty_list()
        # Use first 5 for quick test
        test_faculty = faculty_list[:5]
        tracker.config['faculty']['manual_list'] = test_faculty
        tracker.config['faculty']['auto_fetch_from_website'] = False
    
    # Run the search
    results = tracker.run_search()
    
    print(f"\n🎉 Enhanced search completed successfully!")
    print(f"📁 Check the generated files for your results.")

if __name__ == "__main__":
    main()
