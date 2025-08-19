#!/usr/bin/env python3
"""
CSRR Automated Faculty Media Tracker
A user-friendly tool for fellows to automatically track faculty media appearances

Author: AI Assistant
Date: 2025
"""\ert0-\poiu 

import pandas as pd
from docx import Document
import time
import random
from datetime import datetime, timedelta
import requests
from bs4 import BeautifulSoup
import urllib.parse
import re
import json
import os
from typing import List, Dict, Optional
import argparse
import yaml
from pathlib import Path

class AutomatedFacultyMediaTracker:
    """Automated tool for tracking CSRR faculty media appearances"""
    
    def __init__(self, config_file: str = "config.yaml"):
        self.config = self.load_config(config_file)
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        })
        self.results = []
        
    def load_config(self, config_file: str) -> Dict:
        """Load configuration from YAML file"""
        default_config = {
            'search_period': {
                'start_date': '2025-06-01',
                'end_date': '2025-07-31'
            },
            'output': {
                'excel_filename': 'CSRR_Faculty_Media_Report.xlsx',
                'word_filename': 'CSRR_Faculty_Op-Eds.docx',
                'include_snippets': True,
                'max_results_per_faculty': 5
            },
            'search': {
                'max_results_per_query': 3,
                'delay_between_searches': 3,
                'trusted_sources_only': False
            },
            'faculty': {
                'auto_fetch_from_website': True,
                'manual_list': []  # Fallback list if website fetch fails
            }
        }
        
        if os.path.exists(config_file):
            with open(config_file, 'r') as f:
                user_config = yaml.safe_load(f) or {}
                # Merge user config with defaults
                for key, value in user_config.items():
                    if key in default_config:
                        if isinstance(value, dict):
                            default_config[key].update(value)
                        else:
                            default_config[key] = value
                    else:
                        default_config[key] = value
        
        return default_config
    
    def create_default_config(self, config_file: str = "config.yaml"):
        """Create a default configuration file"""
        default_config = {
            'search_period': {
                'start_date': '2025-06-01',
                'end_date': '2025-07-31',
                'description': 'Date range for media search (YYYY-MM-DD format)'
            },
            'output': {
                'excel_filename': 'CSRR_Faculty_Media_Report.xlsx',
                'word_filename': 'CSRR_Faculty_Media_Report.docx',
                'include_snippets': True,
                'max_results_per_faculty': 5,
                'description': 'Output file settings'
            },
            'search': {
                'max_results_per_query': 3,
                'delay_between_searches': 3,
                'trusted_sources_only': False,
                'description': 'Search behavior settings'
            },
            'faculty': {
                'auto_fetch_from_website': True,
                'manual_list': [],
                'description': 'Faculty list settings'
            }
        }
        
        with open(config_file, 'w') as f:
            yaml.dump(default_config, f, default_flow_style=False, indent=2)
        
        print(f"✅ Created default configuration file: {config_file}")
        print("📝 Edit this file to customize your search settings")
    
    def fetch_faculty_list(self) -> List[str]:
        """Fetch faculty list from CSRR website or use manual list"""
        if self.config['faculty']['auto_fetch_from_website']:
            try:
                print("🌐 Fetching faculty list from CSRR website...")
                faculty_list = self.scrape_faculty_from_website()
                if faculty_list:
                    print(f"✅ Found {len(faculty_list)} faculty members from website")
                    return faculty_list
            except Exception as e:
                print(f"⚠️  Could not fetch from website: {e}")
        
        # Fallback to manual list
        manual_list = self.config['faculty']['manual_list']
        if manual_list:
            print(f"📋 Using manual faculty list: {len(manual_list)} members")
            return manual_list
        
        # Final fallback - use the comprehensive list from the original script
        fallback_list = [
            "Adil Haque", "Adnan Zulfiqar", "Alexander A. Reinert", "Alexander Hinton",
            "Alexis Karteron", "Ali A. Olomi", "Ali R. Chaudhary", "Ameena Ghaffar-Kucher",
            "Amir Hussain", "Anju Gupta", "Atalia Omer", "Atiya Aftab", "Audrey Truschke",
            "Aziz Rana", "Aziza Ahmed", "Behrooz Ghamari-Tabrizi", "Beth Baron", "Beth Stephens",
            "Bidisha Biswas", "Brittany Friedman", "Catherine M. Grosso", "Chaumtoli Huq",
            "Clark B. Lombardi", "Cyra A. Choudhury", "D. Asher Ghertner", "Dalia Fahmy",
            "Deepa Kumar", "Eid Mohamed", "Elise Boddie", "Ellen C. Yaroshefsky",
            "Emily Berman", "Emmaia Gelman", "Eric McDaniel", "Esaa Mohammad Samarah Samarsky",
            "Esther Canty-Barnes", "Faisal Kutty", "Faiza Sayed", "Falguni A. Sheth",
            "Farid Hafez", "Fatemeh Shams", "Gaiutra Devi Bahadur", "Ghada Ageel",
            "Hadi Khoshneviss", "Haider Ala Hamoudi", "Hajar Yazdiha", "Hatem Bazian",
            "Heba M. Khalil", "Huda J. Fakhreddine", "Irus Braverman", "Ivan Kalmar",
            "James R. Jones", "Jasbir K. Puar", "Jasmin Zine", "Jason Brownlee",
            "Jesse Norris", "Joel Beinin", "John Tehranian", "Jon Dubin", "Jonathan Feingold",
            "Jonathan Hafetz", "Jorge Contesse", "Juan Cole", "Karam Dana", "Karim Malak",
            "Karishma Desai", "Khaled A. Beydoun", "Khyati Joshi", "LaToya Baldwin Clark",
            "Lara Sheehi", "Leila Kawar", "Leyla Amzi-Erdogdular", "Lorenzo Veracini",
            "Mahruq Khan", "Marc O. Jones", "Margaret Hu", "Mark Bray", "Marta Esquilin",
            "Maryam Jamshidi", "Matthew Abraham", "Maura Finkelstein", "Mayte Green-Mercado",
            "Meera E. Deo", "Mitra Rastegar", "Mohammad Fadel", "Mojtaba Mahdavi",
            "Muhannad Ayyash", "Nader Hashemi", "Nadia Ahmad", "Nahed Samour",
            "Nancy A. Khalil", "Natsu Taylor Saito", "Nausheen Husain", "Naved Bakali",
            "Nazia Kazi", "Nermin Allam", "Norrinda Brown Hayat", "Noura Erakat",
            "Omar Al-Dewachi", "Omar Dajani", "Omar S. Dahi", "Omid Safi", "Ousseina Alidou",
            "Rabea Benhalim", "Rachel Godsil", "Raquel E Aldana", "Raz Segal",
            "Rebecca Hankins", "Riaz Tejani", "Robert S. Chang", "Sabreena Ghaffar-Siddiqui",
            "Sabrina Alimahomed", "Saeed Khan", "Sahar Mohamed Khamis", "Salman Sayyid",
            "Santiago Slabodsky", "Sarah Eltantawi", "Seema Saifee", "Sherene Razack",
            "Shirin Saeidi", "Shirin Sinnar", "Shoba Sivaprasad Wadhia", "Siham Elkassem",
            "Stacy Hawkins", "Stephen Dycus", "Stephen Sheehi", "Susan M. Akram",
            "Sylvia Chan-Malik", "Tahseen Shah", "Taleed El-Sabawi", "Tazeen M. Ali",
            "Timothy Eatman", "Timothy P. Daniels", "Timothy Raphael", "Toby Jones",
            "Udi Ofer", "Umayyah Cable", "Veena Dubal", "Victoria Ramenzoni",
            "Wadie Said", "Wendell Marsh", "Wendy Greene", "Whitney Strub",
            "William C. Banks", "William I. Robinson", "Yasmiyn Irizarry", "Zahra Ali",
            "Zain Abdullah", "Zakia Salime", "M. Shahid Alam", "Khalil Al-Anani",
            "Joseph Massad", "John L. Esposito", "Katherine M. Franke", "Tanya K. Hernandez"
        ]
        print(f"📋 Using fallback faculty list: {len(fallback_list)} members")
        return fallback_list
    
    def scrape_faculty_from_website(self) -> List[str]:
        """Scrape faculty names from CSRR website"""
        url = 'https://csrr.rutgers.edu/about/faculty-affiliates/'
        response = self.session.get(url)
        soup = BeautifulSoup(response.content, 'html.parser')
        
        faculty_names = []
        text_content = soup.get_text()
        lines = text_content.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            # Look for names that precede academic titles
            next_line = lines[i+1].strip() if i+1 < len(lines) else ""
            
            if any(title in next_line.lower() for title in [
                'professor', 'associate professor', 'assistant professor', 
                'clinical professor', 'emeritus professor', 'visiting professor',
                'distinguished professor', 'lecturer', 'instructor'
            ]):
                name = line.strip()
                name = re.sub(r'^(Dr\.|Prof\.|Professor)\s+', '', name)
                name = re.sub(r'\s+(Jr\.|Sr\.|II|III|IV)$', r' \1', name)
                
                if len(name.split()) >= 2 and len(name) < 50:
                    faculty_names.append(name)
        
        # Also look for names in HTML elements
        for element in soup.find_all(['h3', 'h4', 'h5', 'strong', 'b']):
            text = element.get_text().strip()
            if text and len(text.split()) >= 2 and len(text) < 50:
                words = text.split()
                if all(word[0].isupper() for word in words if word):
                    faculty_names.append(text)
        
        # Clean and filter
        faculty_names = sorted(list(set(faculty_names)))
        filtered_faculty = []
        
        for name in faculty_names:
            if not any(term in name.upper() for term in ['NEWS', 'UPDATE', 'ALERT', 'CLICK', 'VIEW', 'READ']):
                name_parts = name.split()
                if len(name_parts) >= 2 and all(part[0].isupper() and part[1:].islower() for part in name_parts if part.isalpha()):
                    filtered_faculty.append(name)
        
        return filtered_faculty
    
    def search_faculty_media(self, faculty_name: str) -> List[Dict]:
        """Search for media appearances by a faculty member"""
        print(f"🔍 Searching for: {faculty_name}")
        
        search_queries = [
            f'"{faculty_name}" article 2025',
            f'"{faculty_name}" op-ed opinion 2025',
            f'"{faculty_name}" interview 2025',
            f'"{faculty_name}" commentary analysis 2025'
        ]
        
        all_results = []
        
        for query in search_queries:
            try:
                results = self.search_web(query, faculty_name)
                all_results.extend(results)
                
                # Rate limiting
                time.sleep(random.uniform(1, self.config['search']['delay_between_searches']))
                
            except Exception as e:
                print(f"  ⚠️  Search error: {e}")
                continue
        
        # Remove duplicates and limit results
        unique_results = []
        seen_urls = set()
        
        for result in all_results:
            url = result.get('url', '')
            if url and url not in seen_urls:
                seen_urls.add(url)
                unique_results.append(result)
                
                if len(unique_results) >= self.config['output']['max_results_per_faculty']:
                    break
        
        if unique_results:
            print(f"  ✅ Found {len(unique_results)} articles")
        else:
            print(f"  ❌ No articles found")
        
        return unique_results
    
    def search_web(self, query: str, faculty_name: str) -> List[Dict]:
        """Search the web for articles"""
        try:
            # Add date range to query
            start_date = self.config['search_period']['start_date']
            end_date = self.config['search_period']['end_date']
            date_query = f"{query} after:{start_date} before:{end_date}"
            
            # Use Bing search
            url = f"https://www.bing.com/search?q={urllib.parse.quote(date_query)}&count={self.config['search']['max_results_per_query']}"
            
            response = self.session.get(url, timeout=15)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            results = []
            
            for result in soup.find_all('li', class_='b_algo')[:self.config['search']['max_results_per_query']]:
                title_elem = result.find('h2')
                if not title_elem:
                    continue
                    
                title_link = title_elem.find('a')
                if not title_link:
                    continue
                
                title = title_link.get_text(strip=True)
                link = title_link.get('href', '')
                
                snippet_elem = result.find('p')
                snippet = snippet_elem.get_text(strip=True) if snippet_elem else ''
                
                # Validate faculty mention
                if not self.validate_faculty_mention(faculty_name, title, snippet):
                    continue
                
                # Extract source and date
                source = self.extract_source(link)
                pub_date = self.extract_date(f"{title} {snippet}", link)
                
                results.append({
                    'faculty_name': faculty_name,
                    'title': title,
                    'url': link,
                    'snippet': snippet,
                    'source': source,
                    'publication_date': pub_date
                })
            
            return results
            
        except Exception as e:
            print(f"  ⚠️  Web search error: {e}")
            return []
    
    def validate_faculty_mention(self, faculty_name: str, title: str, snippet: str) -> bool:
        """Validate that the faculty member is actually mentioned"""
        content_text = f"{title} {snippet}".lower()
        faculty_lower = faculty_name.lower()
        
        # Check for exact name match
        if faculty_lower in content_text:
            return True
        
        # Check for name components
        name_parts = faculty_lower.split()
        if len(name_parts) >= 2:
            last_name = name_parts[-1]
            first_name = name_parts[0]
            
            if last_name in content_text and first_name in content_text:
                title_words = title.lower().split()
                snippet_words = snippet.lower().split()
                all_words = title_words + snippet_words
                
                try:
                    first_pos = all_words.index(first_name)
                    last_pos = all_words.index(last_name)
                    if abs(first_pos - last_pos) <= 5:
                        return True
                except ValueError:
                    pass
        
        return False
    
    def extract_source(self, url: str) -> str:
        """Extract source name from URL"""
        try:
            domain = urllib.parse.urlparse(url).netloc
            source = domain.replace('www.', '').replace('.com', '').replace('.org', '').replace('.net', '').title()
            return source
        except:
            return "Unknown"
    
    def extract_date(self, content: str, url: str) -> str:
        """Extract publication date"""
        date_patterns = [
            r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+202[45]',
            r'\d{1,2}[/-]\d{1,2}[/-]202[45]',
            r'202[45]-\d{2}-\d{2}'
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                return match.group(0)
        
        return "2025"
    
    def create_excel_report(self, results: List[Dict]) -> str:
        """Create Excel report"""
        if not results:
            df = pd.DataFrame(columns=[
                "Faculty Name", "Title", "Source", "URL", 
                "Publication Date", "Date Found", "Snippet"
            ])
        else:
            excel_data = []
            for result in results:
                excel_data.append({
                    "Faculty Name": result['faculty_name'],
                    "Title": result['title'],
                    "Source": result['source'],
                    "URL": result['url'],
                    "Publication Date": result['publication_date'],
                    "Date Found": datetime.now().strftime('%Y-%m-%d'),
                    "Snippet": result['snippet'][:500] + "..." if len(result['snippet']) > 500 else result['snippet']
                })
            df = pd.DataFrame(excel_data)
        
        filename = self.config['output']['excel_filename']
        df.to_excel(filename, index=False)
        print(f"📊 Excel report saved: {filename}")
        return filename
    
    def create_word_report(self, results: List[Dict]) -> str:
        """Create Word report matching the exact format from Downloads folder"""
        doc = Document()
        
        # Set document properties to match reference
        doc.core_properties.title = "Op-Eds by CSRR Faculty Affiliates"
        doc.core_properties.author = "CSRR Automated Tracker"
        
        # Main title - matches reference format exactly
        title_para = doc.add_heading("Op-Eds by CSRR Faculty Affiliates", 0)
        
        # Subtitle with date range - matches reference format
        start_date = self.config['search_period']['start_date']
        end_date = self.config['search_period']['end_date']
        
        # Convert dates to readable format
        from datetime import datetime
        try:
            start_dt = datetime.strptime(start_date, '%Y-%m-%d')
            end_dt = datetime.strptime(end_date, '%Y-%m-%d')
            
            # Format like "Since October 1, 2023" or "June 1, 2025 - July 31, 2025"
            if start_dt.year == end_dt.year and start_dt.month == 1 and start_dt.day == 1:
                # Full year format
                period = f"Since {start_dt.strftime('%B %-d, %Y')}"
            else:
                # Date range format
                period = f"{start_dt.strftime('%B %-d, %Y')} - {end_dt.strftime('%B %-d, %Y')}"
        except:
            period = f"{start_date} to {end_date}"
        
        subtitle_para = doc.add_paragraph(period)
        doc.add_paragraph("")  # Empty line for spacing
        
        # Group results by faculty
        faculty_results = {}
        for result in results:
            faculty = result['faculty_name']
            if faculty not in faculty_results:
                faculty_results[faculty] = []
            faculty_results[faculty].append(result)
        
        # Add results by faculty - matches reference format exactly
        for faculty_name in sorted(faculty_results.keys()):
            # Add faculty name as a paragraph (not heading) - matches reference format
            doc.add_paragraph(faculty_name)
            
            results = faculty_results[faculty_name]
            for result in results:
                # Format: Author, Title, Source, Date, URL.
                # Note: In the reference, the author is the faculty member
                formatted_entry = (
                    f"{result['faculty_name']}, "
                    f"{result['title']}, "
                    f"{result['source']}, "
                    f"{result['publication_date']}, "
                    f"{result['url']}."
                )
                doc.add_paragraph(formatted_entry)
        
        # Add faculty with no results for completeness
        all_faculty = set(self.fetch_faculty_list())
        faculty_with_results = set(faculty_results.keys())
        faculty_without_results = all_faculty - faculty_with_results
        
        for faculty_name in sorted(faculty_without_results):
            doc.add_paragraph(faculty_name)
            # No articles found - leave empty (matches reference format)
        
        # Create filename with date range for better organization
        start_date = self.config['search_period']['start_date']
        end_date = self.config['search_period']['end_date']
        
        # Convert to readable format for filename
        from datetime import datetime
        try:
            start_dt = datetime.strptime(start_date, '%Y-%m-%d')
            end_dt = datetime.strptime(end_date, '%Y-%m-%d')
            
            if start_dt.year == end_dt.year and start_dt.month == 1 and start_dt.day == 1:
                # Full year format
                date_suffix = f"{start_dt.year}"
            else:
                # Date range format
                date_suffix = f"{start_dt.strftime('%b%Y')}_to_{end_dt.strftime('%b%Y')}"
        except:
            date_suffix = f"{start_date.replace('-', '')}_to_{end_date.replace('-', '')}"
        
        filename = f"CSRR_Faculty_Op-Eds_{date_suffix}.docx"
        doc.save(filename)
        print(f"📄 Word report saved: {filename}")
        return filename
    
    def run_search(self) -> Dict[str, str]:
        """Run the complete media search"""
        print("=" * 60)
        print("🎯 CSRR AUTOMATED FACULTY MEDIA TRACKER")
        print("=" * 60)
        
        # Load faculty list
        faculty_list = self.fetch_faculty_list()
        print(f"👥 Processing {len(faculty_list)} faculty members")
        print(f"📅 Period: {self.config['search_period']['start_date']} to {self.config['search_period']['end_date']}")
        print("=" * 60)
        print()
        
        all_results = []
        faculty_with_results = 0
        
        # Process each faculty member
        for i, faculty_name in enumerate(faculty_list, 1):
            print(f"[{i:3d}/{len(faculty_list)}] ", end="")
            
            results = self.search_faculty_media(faculty_name)
            
            if results:
                faculty_with_results += 1
                all_results.extend(results)
            
            # Progress update
            if i % 20 == 0:
                print(f"\n📊 Progress: {i}/{len(faculty_list)} faculty processed")
                print(f"   Found articles for: {faculty_with_results} faculty")
                print(f"   Total articles: {len(all_results)}")
                print()
        
        # Final statistics
        print("\n" + "=" * 60)
        print("📋 SEARCH COMPLETED")
        print("=" * 60)
        print(f"Total faculty processed: {len(faculty_list)}")
        print(f"Faculty with articles: {faculty_with_results}")
        print(f"Total articles found: {len(all_results)}")
        print(f"Success rate: {faculty_with_results/len(faculty_list)*100:.1f}%")
        
        # Generate reports
        excel_file = self.create_excel_report(all_results)
        word_file = self.create_word_report(all_results)
        
        print(f"\n✅ REPORTS GENERATED:")
        print(f"📊 Excel: {excel_file}")
        print(f"📄 Word: {word_file}")
        
        return {
            'excel': excel_file,
            'word': word_file,
            'total_articles': len(all_results),
            'faculty_with_results': faculty_with_results
        }

def main():
    """Main function with command line interface"""
    parser = argparse.ArgumentParser(description='CSRR Automated Faculty Media Tracker')
    parser.add_argument('--config', default='config.yaml', help='Configuration file path')
    parser.add_argument('--create-config', action='store_true', help='Create a default configuration file')
    parser.add_argument('--quick-test', action='store_true', help='Run a quick test with first 5 faculty')
    
    args = parser.parse_args()
    
    if args.create_config:
        tracker = AutomatedFacultyMediaTracker()
        tracker.create_default_config(args.config)
        return
    
    # Initialize tracker
    tracker = AutomatedFacultyMediaTracker(args.config)
    
    if args.quick_test:
        print("🧪 Running quick test with first 5 faculty members...")
        # Modify config for quick test
        tracker.config['output']['max_results_per_faculty'] = 2
        tracker.config['search']['max_results_per_query'] = 1
        
        faculty_list = tracker.fetch_faculty_list()[:5]
        tracker.config['faculty']['manual_list'] = faculty_list
        tracker.config['faculty']['auto_fetch_from_website'] = False
    
    # Run the search
    results = tracker.run_search()
    
    print(f"\n🎉 Search completed successfully!")
    print(f"📁 Check the generated files for your results.")

if __name__ == "__main__":
    main()
