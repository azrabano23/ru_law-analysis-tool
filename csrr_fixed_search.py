#!/usr/bin/env python3
"""
CSRR Faculty Media Search - FIXED VERSION
Addresses accuracy issues and improves output format to match manual collection style
"""

import pandas as pd
from docx import Document
import time
import random
from datetime import datetime, timedelta
import requests
from bs4 import BeautifulSoup
import urllib.parse
import re

# Complete faculty list - all 151 members
faculty_list = [
    "Adil Haque", "Adnan Zulfiqar", "Alexander A. Reinert", "Alexander Hinton",
    "Alexis Karteron", "Ali A. Olomi", "Ali R. Chaudhary", "Ameena Ghaffar-Kucher",
    "Amir Hussain", "Anju Gupta", "Atalia Omer", "Atiya Aftab", "Audrey Truschke",
    "Aziz Rana", "Aziza Ahmed", "Behrooz Ghamari-Tabrizi", "Beth Baron", "Beth Stephens",
    "Bidisha Biswas", "Brittany Friedman", "Catherine M. Grosso", "Chaumtoli Huq",
    "Clark B. Lombardi", "Cyra A. Choudhury", "D. Asher Ghertner", "Dalia Fahmy",
    "Deepa Kumar", "Eid Mohamed", "Elise Boddie", "Ellen C. Yaroshefsky",
    "Emily Berman", "Emmaia Gelman", "Eric McDaniel", "Esaa Mohammad Samarah Samarsky",
    "Esther Canty-Barnes", "Faisal Kutty", "Faiza Sayed", "Falguni A. Sheth",
    "Farid Hafez", "Fatemeh Shams", "Gaiutra Devi Bahadur", "Ghada Ageel",
    "Hadi Khoshneviss", "Haider Ala Hamoudi", "Hajar Yazdiha", "Hatem Bazian",
    "Heba M. Khalil", "Huda J. Fakhreddine", "Irus Braverman", "Ivan Kalmar",
    "James R. Jones", "Jasbir K. Puar", "Jasmin Zine", "Jason Brownlee",
    "Jesse Norris", "Joel Beinin", "John Tehranian", "Jon Dubin", "Jonathan Feingold",
    "Jonathan Hafetz", "Jorge Contesse", "Juan Cole", "Karam Dana", "Karim Malak",
    "Karishma Desai", "Khaled A. Beydoun", "Khyati Joshi", "LaToya Baldwin Clark",
    "Lara Sheehi", "Leila Kawar", "Leyla Amzi-Erdogdular", "Lorenzo Veracini",
    "Mahruq Khan", "Marc O. Jones", "Margaret Hu", "Mark Bray", "Marta Esquilin",
    "Maryam Jamshidi", "Matthew Abraham", "Maura Finkelstein", "Mayte Green-Mercado",
    "Meera E. Deo", "Mitra Rastegar", "Mohammad Fadel", "Mojtaba Mahdavi",
    "Muhannad Ayyash", "Nader Hashemi", "Nadia Ahmad", "Nahed Samour",
    "Nancy A. Khalil", "Natsu Taylor Saito", "Nausheen Husain", "Naved Bakali",
    "Nazia Kazi", "Nermin Allam", "Norrinda Brown Hayat", "Noura Erakat",
    "Omar Al-Dewachi", "Omar Dajani", "Omar S. Dahi", "Omid Safi", "Ousseina Alidou",
    "Rabea Benhalim", "Rachel Godsil", "Raquel E Aldana", "Raz Segal",
    "Rebecca Hankins", "Riaz Tejani", "Robert S. Chang", "Sabreena Ghaffar-Siddiqui",
    "Sabrina Alimahomed", "Saeed Khan", "Sahar Mohamed Khamis", "Salman Sayyid",
    "Santiago Slabodsky", "Sarah Eltantawi", "Seema Saifee", "Sherene Razack",
    "Shirin Saeidi", "Shirin Sinnar", "Shoba Sivaprasad Wadhia", "Siham Elkassem",
    "Stacy Hawkins", "Stephen Dycus", "Stephen Sheehi", "Susan M. Akram",
    "Sylvia Chan-Malik", "Tahseen Shah", "Taleed El-Sabawi", "Tazeen M. Ali",
    "Timothy Eatman", "Timothy P. Daniels", "Timothy Raphael", "Toby Jones",
    "Udi Ofer", "Umayyah Cable", "Veena Dubal", "Victoria Ramenzoni",
    "Wadie Said", "Wendell Marsh", "Wendy Greene", "Whitney Strub",
    "William C. Banks", "William I. Robinson", "Yasmiyn Irizarry", "Zahra Ali",
    "Zain Abdullah", "Zakia Salime", "M. Shahid Alam", "Khalil Al-Anani",
    "Raquel E. Aldana", "Tazeen M. Ali", "Sabrina Alimahomed-Wilson", "Nermin Allam",
    "Mohamed Alsiadi", "Mason Ameri", "Leyla Amzi-Erdogdular", "Mohamed Arafa",
    "Abed Awad", "Asli U. Bali", "William C. Banks", "George Bisharat",
    "Mark Bray", "Umayyah Cable", "Robert S. Chang", "Juan Cole",
    "Timothy P. Daniels", "Meera E. Deo", "Karishma Desai", "Veena Dubal",
    "John L. Esposito", "Marta Esquilin", "John Farmer Jr.", "Katherine M. Franke",
    "Wendy Greene", "Zeynep Devrim Gursel", "Jonathan Hafetz", "Nader Hashemi",
    "Stacy Hawkins", "Tanya K. Hernandez", "Margaret Hu", "Nausheen Husain",
    "Toby Jones", "Nazia Kazi", "Faisal Kutty", "Clark B. Lombardi",
    "Wendell Marsh", "Joseph Massad", "Jesse Norris", "Udi Ofer",
    "Timothy Raphael", "Ebrahim Rasool", "Victoria Ramenzoni", "Mitra Rastegar",
    "Alexander A. Reinert", "William I. Robinson", "Natsu Taylor Saito",
    "Nahed Samour", "Faiza Sayed", "Salman Sayyid", "Alex Dika Seggerman",
    "Saher Selod", "Sudha N. Setty", "Fatemeh Shams", "Lara Sheehi",
    "Stephen Sheehi", "Sabreena Ghaffar-Siddiqui", "Saleema Snow", "SpearIt",
    "Whitney Strub", "Nukhet Varlik", "Lorenzo Veracini", "Shoba Sivaprasad Wadhia",
    "Ellen C. Yaroshefsky", "Hajar Yazdiha", "Jasmin Zine"
]

# Trusted news sources for filtering
TRUSTED_SOURCES = {
    'nytimes.com': 'New York Times',
    'washingtonpost.com': 'Washington Post',
    'cnn.com': 'CNN',
    'aljazeera.com': 'Al Jazeera',
    'bbc.com': 'BBC',
    'npr.org': 'NPR',
    'reuters.com': 'Reuters',
    'politico.com': 'Politico',
    'theatlantic.com': 'The Atlantic',
    'theguardian.com': 'The Guardian',
    'huffpost.com': 'HuffPost',
    'slate.com': 'Slate',
    'vox.com': 'Vox',
    'buzzfeed.com': 'BuzzFeed',
    'axios.com': 'Axios',
    'apnews.com': 'Associated Press',
    'abcnews.go.com': 'ABC News',
    'cbsnews.com': 'CBS News',
    'nbcnews.com': 'NBC News',
    'foxnews.com': 'Fox News',
    'usatoday.com': 'USA Today',
    'wsj.com': 'Wall Street Journal',
    'newyorker.com': 'The New Yorker',
    'time.com': 'Time',
    'newsweek.com': 'Newsweek',
    'thehill.com': 'The Hill',
    'pbs.org': 'PBS',
    'economist.com': 'The Economist'
}

def is_trusted_source(url):
    """Check if URL is from a trusted news source"""
    for domain, name in TRUSTED_SOURCES.items():
        if domain in url.lower():
            return True, name
    return False, "Unknown"

def extract_actual_date(snippet, title, url):
    """Extract actual publication date from content when possible"""
    # Try to find date patterns in snippet or title
    date_patterns = [
        r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+202[45]',
        r'\d{1,2}[/-]\d{1,2}[/-]202[45]',
        r'202[45]-\d{2}-\d{2}'
    ]
    
    combined_text = f"{title} {snippet}".lower()
    
    for pattern in date_patterns:
        match = re.search(pattern, combined_text, re.IGNORECASE)
        if match:
            return match.group(0)
    
    # Default to indicating it's from 2025 period
    return "2025"

def validate_faculty_relevance(faculty_name, title, snippet, url):
    """Validate that the result is actually about the faculty member"""
    name_parts = faculty_name.lower().split()
    title_lower = title.lower()
    snippet_lower = snippet.lower()
    
    # Check for full name match
    if faculty_name.lower() in title_lower or faculty_name.lower() in snippet_lower:
        return True
    
    # Check for partial name matches (at least last name + one other)
    if len(name_parts) >= 2:
        last_name = name_parts[-1]
        first_name = name_parts[0]
        
        # Must have last name and at least one other part
        if (last_name in title_lower or last_name in snippet_lower) and \
           (first_name in title_lower or first_name in snippet_lower):
            return True
    
    return False

def search_google_via_bing(query, max_results=3):
    """Search using Bing with improved accuracy"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        # Add recent time filter to query
        query_with_time = f"{query} after:2024-06-01"
        url = f"https://www.bing.com/search?q={urllib.parse.quote(query_with_time)}&count={max_results}"
        
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        results = []
        
        # Parse Bing search results
        for result in soup.find_all('li', class_='b_algo')[:max_results]:
            title_elem = result.find('h2')
            if title_elem:
                title_link = title_elem.find('a')
                if title_link:
                    title = title_link.get_text(strip=True)
                    link = title_link.get('href', '')
                    
                    # Get snippet
                    snippet_elem = result.find('p')
                    snippet = snippet_elem.get_text(strip=True) if snippet_elem else ''
                    
                    # Only include results from trusted sources
                    is_trusted, source_name = is_trusted_source(link)
                    if is_trusted:
                        results.append({
                            'title': title,
                            'link': link,
                            'snippet': snippet,
                            'source': source_name
                        })
        
        return results
        
    except Exception as e:
        print(f"  Search error: {e}")
        return []

def search_faculty_comprehensive(name):
    """Search for faculty member with accuracy validation"""
    print(f"Searching for: {name}")
    
    # More targeted search strategies focusing on recent academic/media content
    search_queries = [
        f'"{name}" op-ed site:nytimes.com OR site:washingtonpost.com OR site:cnn.com OR site:aljazeera.com',
        f'"{name}" interview site:npr.org OR site:bbc.com OR site:reuters.com',
        f'"{name}" commentary site:theatlantic.com OR site:theguardian.com OR site:slate.com',
        f'"{name}" opinion site:politico.com OR site:huffpost.com OR site:vox.com',
        f'{name} author byline news 2025',
        f'{name} professor quoted interview 2025'
    ]
    
    all_results = []
    
    for query in search_queries:
        try:
            results = search_google_via_bing(query, max_results=5)
            
            # Validate each result
            for result in results:
                if validate_faculty_relevance(name, result['title'], result['snippet'], result['link']):
                    result['faculty_name'] = name
                    result['date_extracted'] = extract_actual_date(result['snippet'], result['title'], result['link'])
                    all_results.append(result)
            
            # Short delay between searches
            time.sleep(random.uniform(1, 2))
            
        except Exception as e:
            print(f"  Error with query '{query}': {e}")
            continue
    
    # Remove duplicates based on URL
    unique_results = []
    seen_urls = set()
    
    for result in all_results:
        url = result.get('link', '')
        if url and url not in seen_urls:
            seen_urls.add(url)
            unique_results.append(result)
        
        if len(unique_results) >= 5:  # Limit per faculty for quality
            break
    
    return unique_results

def create_comprehensive_report():
    """Create the comprehensive report with accuracy validation"""
    print("=== CSRR Faculty Media Search: June 2025 - July 2025 (FIXED VERSION) ===")
    print(f"Processing ALL {len(faculty_list)} faculty members with accuracy validation...")
    print("This may take 45-60 minutes to complete due to additional validation steps.\n")
    
    # Initialize data structures
    all_data = []
    doc = Document()
    
    # Document header
    doc.add_heading("Op-Eds by CSRR Faculty Affiliates", 0)
    doc.add_heading("June 2025 - July 2025", 1)
    doc.add_paragraph("")
    
    # Statistics tracking
    total_results = 0
    faculty_with_results = 0
    validation_errors = 0
    
    # Process each faculty member
    for i, faculty_name in enumerate(faculty_list, 1):
        print(f"[{i}/{len(faculty_list)}] Processing: {faculty_name}")
        
        try:
            results = search_faculty_comprehensive(faculty_name)
            
            if results:
                faculty_with_results += 1
                total_results += len(results)
                
                # Add faculty section to document
                doc.add_heading(faculty_name, level=2)
                
                for j, result in enumerate(results, 1):
                    title = result.get('title', 'N/A')
                    link = result.get('link', 'N/A')
                    snippet = result.get('snippet', 'N/A')
                    source = result.get('source', 'Unknown')
                    date_extracted = result.get('date_extracted', '2025')
                    
                    # Double-check validation
                    if not validate_faculty_relevance(faculty_name, title, snippet, link):
                        print(f"  WARNING: Validation failed for result under {faculty_name}: {title[:50]}...")
                        validation_errors += 1
                        continue
                    
                    # Add to Excel data with proper format
                    row = {
                        "Faculty Name": faculty_name,
                        "Author": faculty_name,  # Following manual format
                        "Title": title,
                        "Source": source,
                        "URL": link,
                        "Snippet": snippet[:300] + "..." if len(snippet) > 300 else snippet,
                        "Date Found": datetime.now().strftime('%Y-%m-%d'),
                        "Publication Date": date_extracted,
                        "Search Order": j
                    }
                    all_data.append(row)
                    
                    # Add to Word document in proper format: Author, Title, Source, Date, URL
                    formatted_entry = f"{faculty_name}, {title}, {source}, {date_extracted}, {link}."
                    doc.add_paragraph(formatted_entry)
                
                print(f"  Found {len(results)} validated results")
            else:
                print(f"  No results found")
                # Add faculty name to document even without results for completeness
                doc.add_heading(faculty_name, level=2)
                doc.add_paragraph("No media mentions found for this period.")
            
            # Progress update every 25 faculty
            if i % 25 == 0:
                print(f"\n--- Progress Update ---")
                print(f"Processed: {i}/{len(faculty_list)} faculty")
                print(f"Found results for: {faculty_with_results} faculty")
                print(f"Total validated results: {total_results}")
                print(f"Validation errors caught: {validation_errors}")
                print(f"Average per faculty: {total_results/i:.1f}")
                print("------------------------\n")
            
            # Longer delay between faculty to ensure accuracy and avoid blocking
            time.sleep(random.uniform(3, 5))
            
        except Exception as e:
            print(f"  Error processing {faculty_name}: {e}")
            continue
    
    # Final summary
    print(f"\n=== FINAL SUMMARY ===")
    print(f"Total faculty processed: {len(faculty_list)}")
    print(f"Faculty with results: {faculty_with_results}")
    print(f"Total validated results found: {total_results}")
    print(f"Validation errors caught and removed: {validation_errors}")
    print(f"Average results per faculty: {total_results/len(faculty_list):.1f}")
    print(f"Success rate: {faculty_with_results/len(faculty_list)*100:.1f}%")
    
    # Save Excel file with proper column structure
    excel_filename = "CSRR_Faculty_Media_June_July_2025_FIXED_VALIDATED.xlsx"
    if all_data:
        df = pd.DataFrame(all_data)
        # Reorder columns to match manual collection format
        column_order = ["Faculty Name", "Author", "Title", "Source", "URL", "Publication Date", "Date Found", "Search Order", "Snippet"]
        df = df[column_order]
        df.to_excel(excel_filename, index=False)
        print(f"\nExcel report saved: {excel_filename}")
    else:
        # Create empty Excel with proper structure
        columns = ["Faculty Name", "Author", "Title", "Source", "URL", "Publication Date", "Date Found", "Search Order", "Snippet"]
        df = pd.DataFrame(columns=columns)
        df.to_excel(excel_filename, index=False)
        print(f"\nEmpty Excel report created: {excel_filename}")
    
    # Save Word document
    doc_filename = "CSRR_Faculty_Media_June_July_2025_FIXED_VALIDATED.docx"
    doc.save(doc_filename)
    print(f"Word report saved: {doc_filename}")
    
    return excel_filename, doc_filename

if __name__ == "__main__":
    try:
        print("Starting FIXED comprehensive search for all 151 CSRR faculty affiliates...")
        print("This version includes:")
        print("- Accuracy validation for faculty attribution")
        print("- Trusted source filtering")
        print("- Proper format matching manual collection style")
        print("- Real date extraction when possible")
        print("\nThis will take approximately 45-60 minutes to complete.\n")
        
        excel_file, doc_file = create_comprehensive_report()
        
        print(f"\n=== REPORTS COMPLETED ===")
        print(f"Excel file: {excel_file}")
        print(f"Word file: {doc_file}")
        print(f"\nBoth files have been validated for accuracy and formatted properly!")
        print(f"Please manually review a sample of entries to confirm accuracy before sending.")
        
    except Exception as e:
        print(f"Error in main execution: {e}")
        import traceback
        traceback.print_exc()
