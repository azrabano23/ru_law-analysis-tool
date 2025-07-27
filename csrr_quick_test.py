#!/usr/bin/env python3
"""
CSRR Faculty Media Search - QUICK TEST VERSION
Tests the fixes on a smaller sample to demonstrate accuracy improvements
"""

import pandas as pd
from docx import Document
import time
import random
from datetime import datetime
import requests
from bs4 import BeautifulSoup
import urllib.parse
import re

# Test with a smaller subset of high-profile faculty who are likely to have recent media mentions
test_faculty_list = [
    "Ghada Ageel", "Noura Erakat", "Jasbir K. Puar", "Aziz Rana", 
    "Khaled A. Beydoun", "Juan Cole", "Audrey Truschke", "Stephen Sheehi",
    "Lara Sheehi", "Raz Segal", "Deepa Kumar", "Joel Beinin",
    "Marc O. Jones", "Hatem Bazian", "Matthew Abraham", "Joseph Massad"
]

# Trusted news sources for filtering
TRUSTED_SOURCES = {
    'nytimes.com': 'New York Times',
    'washingtonpost.com': 'Washington Post',
    'cnn.com': 'CNN',
    'aljazeera.com': 'Al Jazeera',
    'bbc.com': 'BBC',
    'npr.org': 'NPR',
    'reuters.com': 'Reuters',
    'politico.com': 'Politico',
    'theatlantic.com': 'The Atlantic',
    'theguardian.com': 'The Guardian',
    'huffpost.com': 'HuffPost',
    'slate.com': 'Slate',
    'vox.com': 'Vox',
    'axios.com': 'Axios',
    'apnews.com': 'Associated Press',
    'cbsnews.com': 'CBS News',
    'nbcnews.com': 'NBC News',
    'usatoday.com': 'USA Today',
    'wsj.com': 'Wall Street Journal',
    'newyorker.com': 'The New Yorker',
    'thehill.com': 'The Hill',
    'pbs.org': 'PBS'
}

def is_trusted_source(url):
    """Check if URL is from a trusted news source"""
    for domain, name in TRUSTED_SOURCES.items():
        if domain in url.lower():
            return True, name
    return False, "Unknown"

def validate_faculty_relevance(faculty_name, title, snippet, url):
    """Validate that the result is actually about the faculty member"""
    name_parts = faculty_name.lower().split()
    title_lower = title.lower()
    snippet_lower = snippet.lower()
    
    # Check for full name match
    if faculty_name.lower() in title_lower or faculty_name.lower() in snippet_lower:
        return True
    
    # Check for partial name matches (at least last name + one other)
    if len(name_parts) >= 2:
        last_name = name_parts[-1]
        first_name = name_parts[0]
        
        # Must have last name and at least one other part
        if (last_name in title_lower or last_name in snippet_lower) and \
           (first_name in title_lower or first_name in snippet_lower):
            return True
    
    return False

def search_google_via_bing(query, max_results=3):
    """Search using Bing with improved accuracy"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        url = f"https://www.bing.com/search?q={urllib.parse.quote(query)}&count={max_results}"
        
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        results = []
        
        # Parse Bing search results
        for result in soup.find_all('li', class_='b_algo')[:max_results]:
            title_elem = result.find('h2')
            if title_elem:
                title_link = title_elem.find('a')
                if title_link:
                    title = title_link.get_text(strip=True)
                    link = title_link.get('href', '')
                    
                    # Get snippet
                    snippet_elem = result.find('p')
                    snippet = snippet_elem.get_text(strip=True) if snippet_elem else ''
                    
                    # Only include results from trusted sources
                    is_trusted, source_name = is_trusted_source(link)
                    if is_trusted:
                        results.append({
                            'title': title,
                            'link': link,
                            'snippet': snippet,
                            'source': source_name
                        })
        
        return results
        
    except Exception as e:
        print(f"  Search error: {e}")
        return []

def search_faculty_targeted(name):
    """Search for faculty member with focused, high-quality queries"""
    print(f"Searching for: {name}")
    
    # High-precision search queries focusing on recent content
    search_queries = [
        f'"{name}" 2025 (op-ed OR opinion OR commentary)',
        f'"{name}" Gaza Palestine 2024 2025',
        f'"{name}" interview podcast 2024 2025'
    ]
    
    all_results = []
    
    for query in search_queries:
        try:
            results = search_google_via_bing(query, max_results=3)
            
            # Validate each result
            for result in results:
                if validate_faculty_relevance(name, result['title'], result['snippet'], result['link']):
                    result['faculty_name'] = name
                    all_results.append(result)
            
            # Short delay between searches
            time.sleep(random.uniform(1, 2))
            
        except Exception as e:
            print(f"  Error with query '{query}': {e}")
            continue
    
    # Remove duplicates based on URL
    unique_results = []
    seen_urls = set()
    
    for result in all_results:
        url = result.get('link', '')
        if url and url not in seen_urls:
            seen_urls.add(url)
            unique_results.append(result)
        
        if len(unique_results) >= 3:  # Limit per faculty for quality
            break
    
    return unique_results

def create_test_report():
    """Create a test report with the high-profile faculty subset"""
    print("=== CSRR Faculty Media Search: QUICK TEST VERSION ===")
    print(f"Testing accuracy fixes with {len(test_faculty_list)} high-profile faculty members...")
    print("This should take 10-15 minutes to complete.\n")
    
    # Initialize data structures
    all_data = []
    doc = Document()
    
    # Document header
    doc.add_heading("Op-Eds by CSRR Faculty Affiliates", 0)
    doc.add_heading("June 2025 - July 2025 (TEST VERSION)", 1)
    doc.add_paragraph("")
    
    # Statistics tracking
    total_results = 0
    faculty_with_results = 0
    validation_errors = 0
    
    # Process each faculty member
    for i, faculty_name in enumerate(test_faculty_list, 1):
        print(f"[{i}/{len(test_faculty_list)}] Processing: {faculty_name}")
        
        try:
            results = search_faculty_targeted(faculty_name)
            
            if results:
                faculty_with_results += 1
                total_results += len(results)
                
                # Add faculty section to document
                doc.add_heading(faculty_name, level=2)
                
                for j, result in enumerate(results, 1):
                    title = result.get('title', 'N/A')
                    link = result.get('link', 'N/A')
                    snippet = result.get('snippet', 'N/A')
                    source = result.get('source', 'Unknown')
                    
                    # Double-check validation
                    if not validate_faculty_relevance(faculty_name, title, snippet, link):
                        print(f"  WARNING: Validation failed for result under {faculty_name}: {title[:50]}...")
                        validation_errors += 1
                        continue
                    
                    # Add to Excel data with proper format
                    row = {
                        "Faculty Name": faculty_name,
                        "Author": faculty_name,  # Following manual format
                        "Title": title,
                        "Source": source,
                        "URL": link,
                        "Snippet": snippet[:300] + "..." if len(snippet) > 300 else snippet,
                        "Date Found": datetime.now().strftime('%Y-%m-%d'),
                        "Search Order": j
                    }
                    all_data.append(row)
                    
                    # Add to Word document in proper format: Author, Title, Source, Date, URL
                    formatted_entry = f"{faculty_name}, {title}, {source}, 2025, {link}."
                    doc.add_paragraph(formatted_entry)
                
                print(f"  Found {len(results)} validated results")
            else:
                print(f"  No results found")
                # Add faculty name to document even without results for completeness
                doc.add_heading(faculty_name, level=2)
                doc.add_paragraph("No media mentions found for this test period.")
            
            # Delay between faculty to avoid blocking
            time.sleep(random.uniform(2, 3))
            
        except Exception as e:
            print(f"  Error processing {faculty_name}: {e}")
            continue
    
    # Final summary
    print(f"\n=== TEST RESULTS SUMMARY ===")
    print(f"Total faculty tested: {len(test_faculty_list)}")
    print(f"Faculty with results: {faculty_with_results}")
    print(f"Total validated results found: {total_results}")
    print(f"Validation errors caught and removed: {validation_errors}")
    print(f"Average results per faculty: {total_results/len(test_faculty_list):.1f}")
    print(f"Success rate: {faculty_with_results/len(test_faculty_list)*100:.1f}%")
    
    # Save Excel file with proper column structure
    excel_filename = "CSRR_Faculty_Media_QUICK_TEST_VALIDATED.xlsx"
    if all_data:
        df = pd.DataFrame(all_data)
        # Reorder columns to match manual collection format
        column_order = ["Faculty Name", "Author", "Title", "Source", "URL", "Date Found", "Search Order", "Snippet"]
        df = df[column_order]
        df.to_excel(excel_filename, index=False)
        print(f"\nTest Excel report saved: {excel_filename}")
    else:
        # Create empty Excel with proper structure
        columns = ["Faculty Name", "Author", "Title", "Source", "URL", "Date Found", "Search Order", "Snippet"]
        df = pd.DataFrame(columns=columns)
        df.to_excel(excel_filename, index=False)
        print(f"\nEmpty test Excel report created: {excel_filename}")
    
    # Save Word document
    doc_filename = "CSRR_Faculty_Media_QUICK_TEST_VALIDATED.docx"
    doc.save(doc_filename)
    print(f"Test Word report saved: {doc_filename}")
    
    return excel_filename, doc_filename

if __name__ == "__main__":
    try:
        print("Starting QUICK TEST of accuracy fixes...")
        print("This version demonstrates:")
        print("- Accuracy validation for faculty attribution")
        print("- Trusted source filtering")
        print("- Proper format matching manual collection style")
        print("- Testing with high-profile faculty likely to have recent mentions")
        print("\nThis will take approximately 10-15 minutes to complete.\n")
        
        excel_file, doc_file = create_test_report()
        
        print(f"\n=== TEST REPORTS COMPLETED ===")
        print(f"Excel file: {excel_file}")
        print(f"Word file: {doc_file}")
        print(f"\nNow run the validation script:")
        print(f"python3 validate_results.py {excel_file}")
        
    except Exception as e:
        print(f"Error in test execution: {e}")
        import traceback
        traceback.print_exc()
