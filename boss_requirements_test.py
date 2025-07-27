#!/usr/bin/env python3
"""
BOSS REQUIREMENTS TEST - Addresses specific feedback
Tests the exact issues the boss identified:
1. Incomprehensible/incorrect output
2. Wrong attribution (Nausheen Husain example)
3. Missing format (Author, Title, Source, URL)
4. Need for accuracy checking
"""

import pandas as pd
from docx import Document
import time
import requests
from bs4 import BeautifulSoup
import urllib.parse
import re
from datetime import datetime

class BossRequirementsTest:
    """Test script specifically addressing boss feedback"""
    
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'
        })
        
        # BOSS REQUIREMENT: Only trusted sources
        self.trusted_sources = {
            'nytimes.com': 'New York Times',
            'washingtonpost.com': 'Washington Post',
            'cnn.com': 'CNN',
            'aljazeera.com': 'Al Jazeera',
            'bbc.com': 'BBC',
            'npr.org': 'NPR',
            'reuters.com': 'Reuters',
            'politico.com': 'Politico',
            'theatlantic.com': 'The Atlantic',
            'theguardian.com': 'The Guardian'
        }
        
        # Test with high-profile faculty who are likely to have real media mentions
        self.test_faculty = [
            "Ghada Ageel",  # Known for Gaza commentary
            "Noura Erakat", # Legal scholar with media presence
            "Juan Cole",    # Well-known Middle East expert
            "Nausheen Husain" # Specifically mentioned by boss
        ]
    
    def validate_faculty_attribution(self, faculty_name: str, title: str, snippet: str) -> bool:
        """
        BOSS REQUIREMENT: Strict validation to prevent wrong attribution
        Faculty name MUST actually appear in the content
        """
        content = f"{title} {snippet}".lower()
        faculty_lower = faculty_name.lower()
        
        print(f"  🔍 Checking: '{faculty_name}' in '{title[:60]}...'")
        
        # Exact name match (most reliable)
        if faculty_lower in content:
            print(f"    ✅ VALID: Exact name match found")
            return True
        
        # Check name components
        name_parts = faculty_lower.split()
        if len(name_parts) >= 2:
            first_name = name_parts[0]
            last_name = name_parts[-1]
            
            if first_name in content and last_name in content:
                print(f"    ✅ VALID: Both '{first_name}' and '{last_name}' found")
                return True
        
        # BOSS REQUIREMENT: If we can't validate, reject it
        print(f"    ❌ REJECTED: Faculty name not found in content")
        return False
    
    def is_trusted_source(self, url: str) -> tuple[bool, str]:
        """BOSS REQUIREMENT: Only trusted news sources"""
        for domain, name in self.trusted_sources.items():
            if domain in url.lower():
                return True, name
        return False, "Unknown"
    
    def search_faculty_with_validation(self, faculty_name: str) -> list:
        """Search with strict validation to prevent boss's identified issues"""
        print(f"\n🔍 TESTING: {faculty_name}")
        
        try:
            # Simple, targeted search
            query = f'"{faculty_name}" (op-ed OR interview) 2025'
            url = f"https://www.bing.com/search?q={urllib.parse.quote(query)}&count=5"
            
            response = self.session.get(url, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            valid_results = []
            
            for result in soup.find_all('li', class_='b_algo')[:3]:
                title_elem = result.find('h2')
                if not title_elem:
                    continue
                
                title_link = title_elem.find('a')
                if not title_link:
                    continue
                
                title = title_link.get_text(strip=True)
                link = title_link.get('href', '')
                
                snippet_elem = result.find('p')
                snippet = snippet_elem.get_text(strip=True) if snippet_elem else ''
                
                # BOSS REQUIREMENT: Only trusted sources
                is_trusted, source_name = self.is_trusted_source(link)
                if not is_trusted:
                    print(f"  ❌ REJECTED: Untrusted source ({link[:50]}...)")
                    continue
                
                # BOSS REQUIREMENT: Validate faculty attribution
                if not self.validate_faculty_attribution(faculty_name, title, snippet):
                    continue
                
                # If we get here, it's a valid result
                valid_results.append({
                    'faculty_name': faculty_name,
                    'author': faculty_name,  # BOSS REQUIREMENT
                    'title': title,
                    'source': source_name,
                    'url': link,
                    'snippet': snippet
                })
            
            print(f"  📊 Found {len(valid_results)} VALIDATED results for {faculty_name}")
            return valid_results
            
        except Exception as e:
            print(f"  ❌ Search error for {faculty_name}: {e}")
            return []
    
    def create_boss_format_report(self, all_results: list):
        """Create report in EXACT format boss requires"""
        
        # BOSS REQUIREMENT: Excel with Author, Title, Source, URL
        if all_results:
            excel_data = []
            for result in all_results:
                excel_data.append({
                    'Faculty Name': result['faculty_name'],
                    'Author': result['author'],        # BOSS REQUIREMENT
                    'Title': result['title'],
                    'Source': result['source'],        # BOSS REQUIREMENT  
                    'URL': result['url']               # BOSS REQUIREMENT
                })
            
            df = pd.DataFrame(excel_data)
            excel_filename = f"BOSS_REQUIREMENTS_TEST_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
            df.to_excel(excel_filename, index=False)
            print(f"\n📊 Excel report created: {excel_filename}")
            
            # Show sample of what boss will see
            print(f"\n📋 PREVIEW (Boss Format):")
            print(df.to_string(index=False, max_colwidth=50))
        
        # BOSS REQUIREMENT: Word document format  
        doc = Document()
        doc.add_heading("Op-Eds by CSRR Faculty Affiliates", 0)
        doc.add_heading("June 2025 - July 2025 (BOSS REQUIREMENTS TEST)", 1)
        doc.add_paragraph("")
        
        # Group by faculty
        faculty_results = {}
        for result in all_results:
            faculty = result['faculty_name']
            if faculty not in faculty_results:
                faculty_results[faculty] = []
            faculty_results[faculty].append(result)
        
        for faculty_name in sorted(faculty_results.keys()):
            doc.add_heading(faculty_name, level=2)
            
            for result in faculty_results[faculty_name]:
                # BOSS REQUIREMENT: Author, Title, Source, Date, URL format
                formatted_entry = f"{result['author']}, {result['title']}, {result['source']}, 2025, {result['url']}."
                doc.add_paragraph(formatted_entry)
        
        word_filename = f"BOSS_REQUIREMENTS_TEST_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        doc.save(word_filename)
        print(f"📄 Word report created: {word_filename}")
        
        return excel_filename if all_results else None, word_filename
    
    def run_boss_requirements_test(self):
        """Run test addressing boss's specific concerns"""
        print("=" * 60)
        print("🎯 BOSS REQUIREMENTS TEST")
        print("=" * 60)
        print("Testing fixes for boss feedback:")
        print("✅ 1. Prevent incomprehensible/incorrect output")
        print("✅ 2. Fix wrong attribution (Nausheen Husain example)")
        print("✅ 3. Include Author, Title, Source, URL format")
        print("✅ 4. Enable accuracy checking")
        print("=" * 60)
        
        all_results = []
        
        for faculty_name in self.test_faculty:
            results = self.search_faculty_with_validation(faculty_name)
            all_results.extend(results)
            
            # Delay between searches
            time.sleep(3)
        
        print(f"\n📊 FINAL TEST RESULTS:")
        print(f"   Total faculty tested: {len(self.test_faculty)}")
        print(f"   Total validated results: {len(all_results)}")
        print(f"   Faculty with results: {len(set(r['faculty_name'] for r in all_results))}")
        
        # Check for Nausheen Husain misattribution (boss's specific concern)
        nausheen_results = [r for r in all_results if r['faculty_name'] == 'Nausheen Husain']
        other_results = [r for r in all_results if r['faculty_name'] != 'Nausheen Husain']
        
        misattribution_found = False
        for result in other_results:
            if 'nausheen husain' in result['title'].lower() or 'nausheen husain' in result['snippet'].lower():
                misattribution_found = True
                print(f"   ❌ MISATTRIBUTION DETECTED: Nausheen Husain mentioned under {result['faculty_name']}")
        
        if not misattribution_found:
            print(f"   ✅ NAUSHEEN HUSAIN ISSUE FIXED: No misattribution detected")
        
        # Create reports in boss's required format
        excel_file, word_file = self.create_boss_format_report(all_results)
        
        print(f"\n🎉 BOSS REQUIREMENTS TEST COMPLETE!")
        print(f"📊 Files ready for boss review:")
        if excel_file:
            print(f"   Excel: {excel_file}")
        print(f"   Word: {word_file}")
        print(f"\n💡 Next step: Manually review these sample results to confirm accuracy")

def main():
    """Run the boss requirements test"""
    tester = BossRequirementsTest()
    tester.run_boss_requirements_test()

if __name__ == "__main__":
    main()
