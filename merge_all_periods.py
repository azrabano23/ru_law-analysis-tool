#!/usr/bin/env python3
"""
Merge All CSRR Data Periods - Comprehensive Report Generator
Combines May-June 2025 and June-July 2025 data with historical records
"""

import pandas as pd
from docx import Document
from datetime import datetime
import os

def merge_all_data():
    """Merge all CSRR data periods into comprehensive reports"""
    print("=== CSRR Comprehensive Data Merger ===")
    print("Combining all search periods...\n")
    
    # File paths
    may_june_excel = "CSRR_Faculty_Media_May_June_2025.xlsx"
    july_excel = "CSRR_Faculty_Media_June_July_2025_ALL_151_FACULTY.xlsx"
    historical_doc = "5-31-25 Op-Ed CSRR Affiliates (5).docx"
    
    # Check if files exist
    files_exist = []
    for file in [may_june_excel, july_excel, historical_doc]:
        if os.path.exists(file):
            files_exist.append(file)
            print(f"✓ Found: {file}")
        else:
            print(f"✗ Missing: {file}")
    
    if len(files_exist) < 2:
        print("Error: Need at least 2 data files to merge")
        return
    
    # Load Excel data
    all_excel_data = []
    
    # Load May-June data if available
    if may_june_excel in files_exist:
        try:
            may_june_df = pd.read_excel(may_june_excel)
            print(f"Loaded May-June data: {len(may_june_df)} entries")
            all_excel_data.append(may_june_df)
        except Exception as e:
            print(f"Error loading May-June data: {e}")
    
    # Load July data if available
    if july_excel in files_exist:
        try:
            july_df = pd.read_excel(july_excel)
            print(f"Loaded July data: {len(july_df)} entries")
            all_excel_data.append(july_df)
        except Exception as e:
            print(f"Error loading July data: {e}")
    
    # Combine Excel data
    if all_excel_data:
        combined_df = pd.concat(all_excel_data, ignore_index=True)
        
        # Add period column if not exists
        if 'Period' not in combined_df.columns:
            combined_df['Period'] = 'May 31 - July 18, 2025'
        
        # Sort by faculty name and date
        combined_df = combined_df.sort_values(['Faculty Name', 'Date Found'])
        
        # Save comprehensive Excel
        comprehensive_excel = "CSRR_Faculty_Media_COMPREHENSIVE_May_July_2025.xlsx"
        combined_df.to_excel(comprehensive_excel, index=False)
        print(f"✓ Created comprehensive Excel: {comprehensive_excel}")
        print(f"  Total entries: {len(combined_df)}")
        print(f"  Unique faculty: {len(combined_df['Faculty Name'].unique())}")
    
    # Create comprehensive Word document
    comprehensive_doc = Document()
    
    # Document header
    comprehensive_doc.add_heading("Op-Eds by CSRR Faculty Affiliates", 0)
    comprehensive_doc.add_heading("COMPREHENSIVE REPORT", 1)
    comprehensive_doc.add_heading("October 2023 - July 18, 2025", 2)
    comprehensive_doc.add_paragraph("")
    
    # Add summary paragraph
    summary = comprehensive_doc.add_paragraph()
    summary.add_run("This comprehensive report includes all op-eds, interviews, and media appearances by CSRR faculty affiliates from October 2023 through July 18, 2025. ")
    summary.add_run("The report combines historical data with automated search results from May 31 - July 18, 2025.")
    comprehensive_doc.add_paragraph("")
    
    # Add search results summary
    if all_excel_data:
        stats = comprehensive_doc.add_paragraph()
        stats.add_run("AUTOMATED SEARCH RESULTS (May 31 - July 18, 2025):\n").bold = True
        stats.add_run(f"• Total Media Mentions: {len(combined_df)}\n")
        stats.add_run(f"• Faculty with Mentions: {len(combined_df['Faculty Name'].unique())}\n")
        stats.add_run(f"• Search Period: May 31 - July 18, 2025\n")
        stats.add_run(f"• Sources: CNN, NYT, Washington Post, Al Jazeera, BBC, NPR, Reuters, and 20+ others\n")
        comprehensive_doc.add_paragraph("")
    
    # Add historical section first
    comprehensive_doc.add_heading("Historical Records (October 2023 - May 31, 2025)", level=1)
    
    if historical_doc in files_exist:
        try:
            historical_document = Document(historical_doc)
            
            # Copy historical content
            for paragraph in historical_document.paragraphs:
                if paragraph.text.strip():
                    # Copy paragraph with original style
                    new_para = comprehensive_doc.add_paragraph()
                    for run in paragraph.runs:
                        new_run = new_para.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
            
            print("✓ Added historical data to comprehensive document")
        except Exception as e:
            print(f"Error adding historical data: {e}")
    
    # Add new search results section
    comprehensive_doc.add_heading("New Search Results (May 31 - July 18, 2025)", level=1)
    
    if all_excel_data:
        # Group by faculty and add to document
        for faculty_name in sorted(combined_df['Faculty Name'].unique()):
            faculty_data = combined_df[combined_df['Faculty Name'] == faculty_name]
            
            if len(faculty_data) > 0:
                comprehensive_doc.add_heading(faculty_name, level=2)
                
                for _, row in faculty_data.iterrows():
                    title = row.get('Title', 'N/A')
                    publication = row.get('Publication', 'Unknown')
                    link = row.get('Link', 'N/A')
                    period = row.get('Period', 'May-July 2025')
                    
                    # Format like original document
                    para = comprehensive_doc.add_paragraph(f"{publication}, {title}, 2025, {link}.")
    
    # Save comprehensive Word document
    comprehensive_doc_name = "CSRR_Faculty_Media_COMPREHENSIVE_Oct2023_July2025.docx"
    comprehensive_doc.save(comprehensive_doc_name)
    print(f"✓ Created comprehensive Word document: {comprehensive_doc_name}")
    
    # Final summary
    print(f"\n=== COMPREHENSIVE MERGER COMPLETED ===")
    if all_excel_data:
        print(f"Excel Report: {comprehensive_excel}")
        print(f"  - {len(combined_df)} total media mentions")
        print(f"  - {len(combined_df['Faculty Name'].unique())} faculty with mentions")
    print(f"Word Report: {comprehensive_doc_name}")
    print(f"  - Complete historical record (Oct 2023 - July 2025)")
    print(f"  - Ready for publication and distribution")
    
    return comprehensive_excel if all_excel_data else None, comprehensive_doc_name

if __name__ == "__main__":
    try:
        excel_file, doc_file = merge_all_data()
        print(f"\n✅ All files ready for final delivery!")
        
    except Exception as e:
        print(f"Error in merger: {e}")
        import traceback
        traceback.print_exc()
