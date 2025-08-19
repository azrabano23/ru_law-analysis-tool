#!/usr/bin/env python3
"""
Analyze the format of the reference Word document
"""

from docx import Document
import os
import sys

def analyze_document_format(filename):
    """Analyze the structure and format of a Word document"""
    if not os.path.exists(filename):
        print(f"❌ File not found: {filename}")
        return
    
    doc = Document(filename)
    
    print("=" * 60)
    print(f"📄 ANALYZING DOCUMENT: {filename}")
    print("=" * 60)
    
    print(f"\n📊 Document Statistics:")
    print(f"   Total paragraphs: {len(doc.paragraphs)}")
    print(f"   Total sections: {len(doc.sections)}")
    
    print(f"\n📋 Document Structure:")
    
    # Analyze headings and structure
    heading_count = 0
    faculty_sections = []
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            # Check if it's a heading
            if para.style.name.startswith('Heading'):
                heading_count += 1
                print(f"   Heading {heading_count}: '{para.text}' (Style: {para.style.name})")
                
                # Check if it looks like a faculty name
                if len(para.text.split()) >= 2 and not any(word.lower() in para.text.lower() for word in ['op-ed', 'faculty', 'affiliates', 'csrr', 'media']):
                    faculty_sections.append(para.text)
            
            # Check for specific patterns
            elif 'op-ed' in para.text.lower() or 'faculty' in para.text.lower():
                print(f"   Special section: '{para.text}'")
    
    print(f"\n👥 Faculty Sections Found: {len(faculty_sections)}")
    for i, faculty in enumerate(faculty_sections[:10], 1):  # Show first 10
        print(f"   {i}. {faculty}")
    
    if len(faculty_sections) > 10:
        print(f"   ... and {len(faculty_sections) - 10} more")
    
    # Analyze the first few paragraphs to understand the format
    print(f"\n📝 Sample Content Analysis:")
    for i, para in enumerate(doc.paragraphs[:20]):
        if para.text.strip():
            print(f"   Para {i+1}: '{para.text[:80]}{'...' if len(para.text) > 80 else ''}'")
            if i >= 10:  # Show first 10 non-empty paragraphs
                break
    
    # Check for specific formatting patterns
    print(f"\n🎨 Formatting Analysis:")
    
    # Look for the format pattern in entries
    entry_patterns = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and ',' in text and len(text.split(',')) >= 4:
            # This might be an article entry
            parts = text.split(',')
            if len(parts) >= 4:
                entry_patterns.append({
                    'text': text,
                    'parts': len(parts),
                    'sample': parts[:4]  # First 4 parts
                })
    
    if entry_patterns:
        print(f"   Article entry format detected:")
        sample = entry_patterns[0]
        print(f"   Format: {sample['parts']} parts separated by commas")
        print(f"   Sample: {sample['text'][:100]}{'...' if len(sample['text']) > 100 else ''}")
        
        # Analyze the parts
        parts = sample['text'].split(',')
        print(f"   Part 1 (Author): {parts[0].strip()}")
        print(f"   Part 2 (Title): {parts[1].strip()}")
        print(f"   Part 3 (Source): {parts[2].strip()}")
        print(f"   Part 4 (Date): {parts[3].strip()}")
        if len(parts) > 4:
            print(f"   Part 5 (URL): {parts[4].strip()}")
    
    # Check document properties
    print(f"\n📋 Document Properties:")
    core_props = doc.core_properties
    if core_props.title:
        print(f"   Title: {core_props.title}")
    if core_props.author:
        print(f"   Author: {core_props.author}")
    if core_props.created:
        print(f"   Created: {core_props.created}")
    if core_props.modified:
        print(f"   Modified: {core_props.modified}")

if __name__ == "__main__":
    # Use command line argument if provided, otherwise use default
    filename = sys.argv[1] if len(sys.argv) > 1 else "reference_format.docx"
    analyze_document_format(filename)
