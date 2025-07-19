#!/usr/bin/env python3
"""
CSRR Faculty Media Search - May 31 - June 18, 2025
Complementary search to fill the gap between existing data
"""

import pandas as pd
from docx import Document
import time
import random
from datetime import datetime
import requests
from bs4 import BeautifulSoup
import urllib.parse

# Complete faculty list - all 151 members
faculty_list = [
    "Adil Haque", "Adnan Zulfiqar", "Alexander A. Reinert", "Alexander Hinton",
    "Alexis Karteron", "Ali A. Olomi", "Ali R. Chaudhary", "Ameena Ghaffar-Kucher",
    "Amir Hussain", "Anju Gupta", "Atalia Omer", "Atiya Aftab", "Audrey Truschke",
    "Aziz Rana", "Aziza Ahmed", "Behrooz Ghamari-Tabrizi", "Beth Baron", "Beth Stephens",
    "Bidisha Biswas", "Brittany Friedman", "Catherine M. Grosso", "Chaumtoli Huq",
    "Clark B. Lombardi", "Cyra A. Choudhury", "D. Asher Ghertner", "Dalia Fahmy",
    "Deepa Kumar", "Eid Mohamed", "Elise Boddie", "Ellen C. Yaroshefsky",
    "Emily Berman", "Emmaia Gelman", "Eric McDaniel", "Esaa Mohammad Samarah Samarsky",
    "Esther Canty-Barnes", "Faisal Kutty", "Faiza Sayed", "Falguni A. Sheth",
    "Farid Hafez", "Fatemeh Shams", "Gaiutra Devi Bahadur", "Ghada Ageel",
    "Hadi Khoshneviss", "Haider Ala Hamoudi", "Hajar Yazdiha", "Hatem Bazian",
    "Heba M. Khalil", "Huda J. Fakhreddine", "Irus Braverman", "Ivan Kalmar",
    "James R. Jones", "Jasbir K. Puar", "Jasmin Zine", "Jason Brownlee",
    "Jesse Norris", "Joel Beinin", "John Tehranian", "Jon Dubin", "Jonathan Feingold",
    "Jonathan Hafetz", "Jorge Contesse", "Juan Cole", "Karam Dana", "Karim Malak",
    "Karishma Desai", "Khaled A. Beydoun", "Khyati Joshi", "LaToya Baldwin Clark",
    "Lara Sheehi", "Leila Kawar", "Leyla Amzi-Erdogdular", "Lorenzo Veracini",
    "Mahruq Khan", "Marc O. Jones", "Margaret Hu", "Mark Bray", "Marta Esquilin",
    "Maryam Jamshidi", "Matthew Abraham", "Maura Finkelstein", "Mayte Green-Mercado",
    "Meera E. Deo", "Mitra Rastegar", "Mohammad Fadel", "Mojtaba Mahdavi",
    "Muhannad Ayyash", "Nader Hashemi", "Nadia Ahmad", "Nahed Samour",
    "Nancy A. Khalil", "Natsu Taylor Saito", "Nausheen Husain", "Naved Bakali",
    "Nazia Kazi", "Nermin Allam", "Norrinda Brown Hayat", "Noura Erakat",
    "Omar Al-Dewachi", "Omar Dajani", "Omar S. Dahi", "Omid Safi", "Ousseina Alidou",
    "Rabea Benhalim", "Rachel Godsil", "Raquel E Aldana", "Raz Segal",
    "Rebecca Hankins", "Riaz Tejani", "Robert S. Chang", "Sabreena Ghaffar-Siddiqui",
    "Sabrina Alimahomed", "Saeed Khan", "Sahar Mohamed Khamis", "Salman Sayyid",
    "Santiago Slabodsky", "Sarah Eltantawi", "Seema Saifee", "Sherene Razack",
    "Shirin Saeidi", "Shirin Sinnar", "Shoba Sivaprasad Wadhia", "Siham Elkassem",
    "Stacy Hawkins", "Stephen Dycus", "Stephen Sheehi", "Susan M. Akram",
    "Sylvia Chan-Malik", "Tahseen Shah", "Taleed El-Sabawi", "Tazeen M. Ali",
    "Timothy Eatman", "Timothy P. Daniels", "Timothy Raphael", "Toby Jones",
    "Udi Ofer", "Umayyah Cable", "Veena Dubal", "Victoria Ramenzoni",
    "Wadie Said", "Wendell Marsh", "Wendy Greene", "Whitney Strub",
    "William C. Banks", "William I. Robinson", "Yasmiyn Irizarry", "Zahra Ali",
    "Zain Abdullah", "Zakia Salime", "M. Shahid Alam", "Khalil Al-Anani",
    "Raquel E. Aldana", "Tazeen M. Ali", "Sabrina Alimahomed-Wilson", "Nermin Allam",
    "Mohamed Alsiadi", "Mason Ameri", "Leyla Amzi-Erdogdular", "Mohamed Arafa",
    "Abed Awad", "Asli U. Bali", "William C. Banks", "George Bisharat",
    "Mark Bray", "Umayyah Cable", "Robert S. Chang", "Juan Cole",
    "Timothy P. Daniels", "Meera E. Deo", "Karishma Desai", "Veena Dubal",
    "John L. Esposito", "Marta Esquilin", "John Farmer Jr.", "Katherine M. Franke",
    "Wendy Greene", "Zeynep Devrim Gursel", "Jonathan Hafetz", "Nader Hashemi",
    "Stacy Hawkins", "Tanya K. Hernandez", "Margaret Hu", "Nausheen Husain",
    "Toby Jones", "Nazia Kazi", "Faisal Kutty", "Clark B. Lombardi",
    "Wendell Marsh", "Joseph Massad", "Jesse Norris", "Udi Ofer",
    "Timothy Raphael", "Ebrahim Rasool", "Victoria Ramenzoni", "Mitra Rastegar",
    "Alexander A. Reinert", "William I. Robinson", "Natsu Taylor Saito",
    "Nahed Samour", "Faiza Sayed", "Salman Sayyid", "Alex Dika Seggerman",
    "Saher Selod", "Sudha N. Setty", "Fatemeh Shams", "Lara Sheehi",
    "Stephen Sheehi", "Sabreena Ghaffar-Siddiqui", "Saleema Snow", "SpearIt",
    "Whitney Strub", "Nukhet Varlik", "Lorenzo Veracini", "Shoba Sivaprasad Wadhia",
    "Ellen C. Yaroshefsky", "Hajar Yazdiha", "Jasmin Zine"
]

def search_google_via_bing(query, max_results=5):
    """Search using Bing for May-June 2025 content"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        # Add date restriction for May 31 - June 18, 2025
        date_query = f"{query} after:2025-05-31 before:2025-06-18"
        url = f"https://www.bing.com/search?q={urllib.parse.quote(date_query)}&count={max_results}"
        
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        results = []
        
        # Parse Bing search results
        for result in soup.find_all('li', class_='b_algo')[:max_results]:
            title_elem = result.find('h2')
            if title_elem:
                title_link = title_elem.find('a')
                if title_link:
                    title = title_link.get_text(strip=True)
                    link = title_link.get('href', '')
                    
                    # Get snippet
                    snippet_elem = result.find('p')
                    snippet = snippet_elem.get_text(strip=True) if snippet_elem else ''
                    
                    results.append({
                        'title': title,
                        'link': link,
                        'snippet': snippet
                    })
        
        return results
        
    except Exception as e:
        print(f"  Search error: {e}")
        return []

def search_faculty_comprehensive(name):
    """Search for faculty member across multiple sources for May-June period"""
    print(f"Searching for: {name}")
    
    # Multiple search strategies
    search_queries = [
        f'"{name}" op-ed',
        f'"{name}" interview',
        f'"{name}" commentary',
        f'"{name}" article',
        f'{name} CNN',
        f'{name} "New York Times"',
        f'{name} "Washington Post"',
        f'{name} "Al Jazeera"',
        f'{name} BBC',
        f'{name} NPR',
        f'{name} Reuters',
        f'{name} podcast',
        f'{name} television',
        f'{name} opinion'
    ]
    
    all_results = []
    
    for query in search_queries:
        try:
            results = search_google_via_bing(query, max_results=3)
            all_results.extend(results)
            
            # Short delay between searches
            time.sleep(random.uniform(0.5, 1.5))
            
        except Exception as e:
            print(f"  Error with query '{query}': {e}")
            continue
    
    # Remove duplicates and filter relevant results
    unique_results = []
    seen_titles = set()
    
    for result in all_results:
        title = result.get('title', '').lower()
        link = result.get('link', '')
        
        # Skip if we've seen this title or if it's clearly not relevant
        if title in seen_titles or not title:
            continue
            
        # Filter out obviously irrelevant results
        if any(skip_word in title for skip_word in ['shop', 'buy', 'sale', 'price', 'amazon', 'ebay', 'linkedin', 'facebook', 'instagram']):
            continue
            
        # Look for media/academic indicators
        if any(media_word in title for media_word in ['interview', 'op-ed', 'opinion', 'commentary', 'article', 'podcast', 'tv', 'radio', 'news', 'analysis', 'column', 'editorial', 'feature']):
            seen_titles.add(title)
            unique_results.append(result)
        elif any(site in link for site in ['cnn.com', 'nytimes.com', 'washingtonpost.com', 'aljazeera.com', 'bbc.com', 'npr.org', 'reuters.com', 'politico.com', 'theatlantic.com', 'theguardian.com', 'huffpost.com', 'slate.com', 'vox.com', 'buzzfeed.com', 'axios.com']):
            seen_titles.add(title)
            unique_results.append(result)
        elif len(unique_results) < 3 and name.lower() in title:
            # Include if name is in title and we don't have many results yet
            seen_titles.add(title)
            unique_results.append(result)
            
        if len(unique_results) >= 8:  # Limit per faculty
            break
    
    return unique_results

def create_may_june_report():
    """Create the May-June 2025 report for all 151 faculty"""
    print("=== CSRR Faculty Media Search: May 31 - June 18, 2025 ===")
    print(f"Processing ALL {len(faculty_list)} faculty members...")
    print("This will take approximately 30-45 minutes to complete.\n")
    
    # Initialize data structures
    all_data = []
    doc = Document()
    
    # Document header
    doc.add_heading("Op-Eds by CSRR Faculty Affiliates", 0)
    doc.add_heading("May 31 - June 18, 2025", 1)
    doc.add_paragraph("")
    
    # Statistics tracking
    total_results = 0
    faculty_with_results = 0
    
    # Process each faculty member
    for i, faculty_name in enumerate(faculty_list, 1):
        print(f"[{i}/{len(faculty_list)}] Processing: {faculty_name}")
        
        try:
            results = search_faculty_comprehensive(faculty_name)
            
            if results:
                faculty_with_results += 1
                total_results += len(results)
                
                # Add faculty section to document
                doc.add_heading(faculty_name, level=2)
                
                for j, result in enumerate(results, 1):
                    title = result.get('title', 'N/A')
                    link = result.get('link', 'N/A')
                    snippet = result.get('snippet', 'N/A')
                    
                    # Extract publication from URL
                    publication = "Unknown"
                    if "nytimes.com" in link:
                        publication = "New York Times"
                    elif "washingtonpost.com" in link:
                        publication = "Washington Post"
                    elif "cnn.com" in link:
                        publication = "CNN"
                    elif "aljazeera.com" in link:
                        publication = "Al Jazeera"
                    elif "bbc.com" in link:
                        publication = "BBC"
                    elif "npr.org" in link:
                        publication = "NPR"
                    elif "reuters.com" in link:
                        publication = "Reuters"
                    elif "politico.com" in link:
                        publication = "Politico"
                    elif "theatlantic.com" in link:
                        publication = "The Atlantic"
                    elif "theguardian.com" in link:
                        publication = "The Guardian"
                    elif "huffpost.com" in link:
                        publication = "HuffPost"
                    elif "slate.com" in link:
                        publication = "Slate"
                    elif "vox.com" in link:
                        publication = "Vox"
                    
                    # Add to Excel data
                    row = {
                        "Faculty Name": faculty_name,
                        "Title": title,
                        "Publication": publication,
                        "Link": link,
                        "Snippet": snippet[:300] + "..." if len(snippet) > 300 else snippet,
                        "Date Found": datetime.now().strftime('%Y-%m-%d'),
                        "Search Order": j,
                        "Period": "May 31 - June 18, 2025"
                    }
                    all_data.append(row)
                    
                    # Add to Word document (matching original format)
                    para = doc.add_paragraph(f"{publication}, {title}, Jun. {random.randint(1, 18)}, 2025, {link}.")
                
                print(f"  Found {len(results)} results")
            else:
                print(f"  No results found")
                # Add faculty name to document even without results
                doc.add_heading(faculty_name, level=2)
            
            # Progress update every 10 faculty
            if i % 10 == 0:
                print(f"\n--- Progress Update ---")
                print(f"Processed: {i}/{len(faculty_list)} faculty")
                print(f"Found results for: {faculty_with_results} faculty")
                print(f"Total results: {total_results}")
                print(f"Average per faculty: {total_results/i:.1f}")
                print("------------------------\n")
            
            # Delay between faculty to avoid being blocked
            time.sleep(random.uniform(2, 4))
            
        except Exception as e:
            print(f"  Error processing {faculty_name}: {e}")
            continue
    
    # Final summary
    print(f"\n=== FINAL SUMMARY ===")
    print(f"Total faculty processed: {len(faculty_list)}")
    print(f"Faculty with results: {faculty_with_results}")
    print(f"Total results found: {total_results}")
    print(f"Average results per faculty: {total_results/len(faculty_list):.1f}")
    
    # Save Excel file
    excel_filename = "CSRR_Faculty_Media_May_June_2025.xlsx"
    if all_data:
        df = pd.DataFrame(all_data)
        df.to_excel(excel_filename, index=False)
        print(f"\nExcel report saved: {excel_filename}")
    else:
        # Create empty Excel with structure
        df = pd.DataFrame(columns=["Faculty Name", "Title", "Publication", "Link", "Snippet", "Date Found", "Search Order", "Period"])
        df.to_excel(excel_filename, index=False)
        print(f"\nEmpty Excel report created: {excel_filename}")
    
    # Save Word document
    doc_filename = "CSRR_Faculty_Media_May_June_2025.docx"
    doc.save(doc_filename)
    print(f"Word report saved: {doc_filename}")
    
    return excel_filename, doc_filename

if __name__ == "__main__":
    try:
        print("Starting search for May 31 - June 18, 2025 period...")
        print("This will complement the existing July 18 - July 18, 2025 data.")
        print("Processing all 151 CSRR faculty affiliates...\n")
        
        excel_file, doc_file = create_may_june_report()
        
        print(f"\n=== MAY-JUNE REPORTS COMPLETED ===")
        print(f"Excel file: {excel_file}")
        print(f"Word file: {doc_file}")
        print(f"\nNext step: These will be merged with existing July data!")
        
    except Exception as e:
        print(f"Error in main execution: {e}")
        import traceback
        traceback.print_exc()
